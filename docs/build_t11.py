# -*- coding: utf-8 -*-
# =============================================================================
# !!! ADVERTENCIA — NO REGENERAR ESTE DOCUMENTO SIN ACTUALIZARLO ANTES !!!
# -----------------------------------------------------------------------------
# El contenido que este script genera describe el negocio como "distribuidora
# mayorista B2B" (ver la seccion 1.1, mas abajo). Ese marco fue REFUTADO por
# docs/estrategico/DIAGNOSTICO_SEGMENTO_CLIENTE.md (2026-07-30), que midio las
# siete dimensiones de compra y cerro con veredicto (c) POBLACION HOMOGENEA:
# no existe segmentacion B2B/B2C en los datos ni forma honesta de derivarla
# (99,94 % de las 10.384 lineas de pedido piden entre 1 y 4 unidades; 0 RUC en
# 3.887 facturas; grupo_cliente / segmento_cliente / cliente_segmento vacias).
#
# El .docx y el .pdf YA ENTREGADOS se conservan TAL CUAL, de forma deliberada:
# son el entregable de su fecha (2026-07-26) y no se reescriben a posteriori.
# Por eso este script tampoco se modifico en su contenido: solo lleva este
# aviso.
#
# ANTES de volver a ejecutarlo hay que, como minimo:
#   1. Cambiar la caracterizacion del negocio a "comercio minorista multicanal
#      de ticket alto" (seccion 1.1 y toda narrativa de jefatura que hable de
#      cliente mayorista, corporativo o de volumen). OJO: la compra POR VOLUMEN
#      al proveedor si es real y se conserva; lo que cambia es el lado de la
#      VENTA.
#   2. Sincronizar los conteos y la factibilidad con la version VIGENTE de
#      docs/tactico/CATALOGO_OBJETIVOS_TACTICOS.md (este script quedo en 68
#      objetivos / 29 simples de la version 3; el catalogo va por la version 5,
#      con 69 objetivos y 30 simples tras incorporarse OTD-VEN-16).
#   3. Revisar docs/estrategico/BASE_ESTRATEGICA.md (version vigente) para los
#      nombres y textos de los objetivos estrategicos, que cambiaron.
# Regenerar sin hacer lo anterior reintroduciria en un entregable un marco de
# negocio que la evidencia ya descarto.
# =============================================================================
"""
Generador del documento RetailMind_T11_Analisis_Tactico.docx (version Word EDITABLE).

Version 3 (2026-07-26): contenido regenerado desde el catalogo SINCRONIZADO
docs/tactico/CATALOGO_OBJETIVOS_TACTICOS.md — 68 objetivos tacticos en SEIS
departamentos, cada uno re-verificado contra la base real retailmind
(110 tablas en el esquema public) via MCP el 2026-07-26.

Cambios de esta version respecto de la 2:
  * Factibilidad recalculada: 66 FACTIBLE HOY / 0 REQUIERE CAMBIO / 2 REQUIERE
    VOLUMEN (antes 44/10/14). Las diez brechas de sistema estan cerradas en sus
    tres capas y la operacion historica de 19 meses esta sembrada.
  * OTD-LOG-11 reclasificado de COMPUESTO a SIMPLE => 29 simples / 39 compuestos.
  * Reencuadre del negocio: distribuidora mayorista B2B (no comercio minorista).
  * Seccion NUEVA de reportes simples implementados (28 informes por pantalla),
    que responde a la segunda parte del enunciado de la tarea.

Salida de esta ejecucion: SOLO el archivo .docx (el PDF lo produce el usuario
tras la revision final).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Estilos base ──────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def _set_heading_color(style_name, size, rgb):
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor(*rgb)
    st.font.bold = True

_set_heading_color("Heading 1", 18, (0x1F, 0x3B, 0x5B))
_set_heading_color("Heading 2", 14, (0x1F, 0x3B, 0x5B))
_set_heading_color("Heading 3", 12, (0x2E, 0x5A, 0x88))

# ── Helpers ────────────────────────────────────────────────────────────────
def h(level, text):
    return doc.add_heading(text, level=level)

def p(text="", bold=False, italic=False, size=None, align=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        par.alignment = align
    return par

def table(headers, rows, widths=None, font_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run("" if val is None else str(val))
            r.font.size = Pt(font_size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t

def merge_first_col(t, rows_data, key_index=0, font_size=9):
    """Fusiona verticalmente las celdas de la primera columna cuando el valor
    (rows_data[i][key_index]) se repite en filas consecutivas."""
    data_rows = t.rows[1:]
    start = 0
    for i in range(1, len(rows_data) + 1):
        if i == len(rows_data) or rows_data[i][key_index] != rows_data[start][key_index]:
            if i - start > 1:
                merged = data_rows[start].cells[0].merge(data_rows[i - 1].cells[0])
                merged.text = ""
                r = merged.paragraphs[0].add_run(rows_data[start][key_index])
                r.bold = True
                r.font.size = Pt(font_size)
            start = i

CENTER = WD_ALIGN_PARAGRAPH.CENTER

# ══════════════ CATÁLOGO DE OBJETIVOS (fuente única del documento) ════════
# Cada tupla: (departamento, id, objetivo, es_simple, soporte_verificado,
#              estado, visualizacion, roles, reporte_implementado)
BDR = "BDR (PostgreSQL)"
BDC = "BDColumnar (ClickHouse)"
FH = "FACTIBLE HOY"
RC = "REQUIERE CAMBIO EN EL SISTEMA"
RV = "REQUIERE VOLUMEN DE DATOS"

VEN = "VENTAS"
COM = "COMPRAS"
INV = "INVENTARIO / BODEGA"
LOG = "LOGÍSTICA / DESPACHO"
SOP = "SOPORTE"
GER = "GERENCIA / DIRECCIÓN (incluye Marketing)"

OBJ = [
    # ── VENTAS ────────────────────────────────────────────────────────────
    (VEN, "OTD-VEN-01",
     "Ver toda la cartera de pedidos y en qué paso del proceso está cada uno hoy "
     "(confirmado, pagado, facturado, en preparación, despachado, entregado).",
     True,
     "pedido.estado_pedido_id → estado_pedido.codigo/nombre (11 estados) + "
     "pedido.numero/fecha_pedido/canal/total (4 083 pedidos reales en 19 meses).",
     FH, "Tabla con filtros por estado y canal", "Vendedor, Gerente, Administrador", True),
    (VEN, "OTD-VEN-02",
     "Controlar cuántos pedidos registra cada vendedor y por cuánto dinero, para "
     "evaluar el cumplimiento individual.",
     True,
     "pedido.vendedor_id (FK a usuario.nombre/apellido; 1 870 pedidos con vendedor y "
     "9 vendedores distintos — NULL solo en pedidos en línea, por diseño) + pedido.total.",
     FH, "Tabla de vendedores ordenada por monto",
     "Gerente, Administrador (el Vendedor ve solo lo propio)", True),
    (VEN, "OTD-VEN-03",
     "Conocer los 10 productos que más se venden — el «producto estrella» — en el "
     "período elegido.",
     False,
     "pedido_detalle.producto_variante_id/cantidad/precio_unitario (10 384 líneas, "
     "834 variantes distintas vendidas) + producto_variante.producto_id → producto.nombre.",
     FH, "Barras de los 10 primeros, con filtro de período",
     "Gerente, Vendedor, Compras, Analista, Administrador", False),
    (VEN, "OTD-VEN-04",
     "Conocer los 10 productos que no se venden o llevan más tiempo sin venderse — el "
     "«producto hueso» — para liquidarlos o dejar de comprarlos.",
     False,
     "producto_variante (1 221 variantes) cruzada contra pedido_detalle."
     "producto_variante_id: 387 variantes sin una sola venta. Última salida por venta en "
     "movimiento_inventario.fecha_creacion + tipo_movimiento.codigo='salida_venta'.",
     FH, "Tabla de rezagados con días sin venta",
     "Gerente, Compras, Analista, Administrador", False),
    (VEN, "OTD-VEN-05",
     "Saber cuánto compra cada cliente: total gastado, número de pedidos y fecha de la "
     "última compra — mirar el negocio desde el cliente, no solo desde la venta.",
     False,
     "pedido.cliente_id/total/fecha_pedido + cliente.nombre/apellido/email — 72 clientes, "
     "69 de ellos con pedidos, repartidos en 19 meses (antes solo 2 cuentas demo).",
     FH, "Tabla de clientes ordenada por monto acumulado",
     "Gerente, Vendedor, Analista, Administrador", False),
    (VEN, "OTD-VEN-06",
     "Ver cómo evolucionan las ventas mes a mes y por categoría de producto.",
     False,
     "pedido.fecha_pedido (19 meses: 2025-01 a 2026-07) + pedido_detalle.cantidad/"
     "precio_unitario/monto_descuento + producto_categoria.categoria_id → categoria.nombre "
     "(11 categorías).",
     FH, "Líneas por mes con desglose por categoría",
     "Gerente, Analista, Administrador", False),
    (VEN, "OTD-VEN-07",
     "Conocer el valor promedio de cada pedido por período y por canal de venta "
     "(mostrador, teléfono, tienda en línea).",
     False,
     "pedido.total/fecha_pedido/canal (4 083 pedidos: 2 213 web, 1 030 mostrador, "
     "840 teléfono).",
     FH, "Tarjetas y línea temporal por canal", "Gerente, Analista, Administrador", False),
    (VEN, "OTD-VEN-08",
     "Detectar los carritos de compra que los clientes dejaron a medias sin llegar a "
     "pagar.",
     True,
     "carrito.estado + COALESCE(fecha_actualizacion, fecha_creacion) + carrito_item."
     "producto_variante_id/cantidad — 290 carritos: 216 abandonados y 54 activos. El "
     "trigger touch no dispara en los abandonados, de ahí el COALESCE para la antigüedad.",
     FH, "Tabla de carritos inactivos con antigüedad",
     "Gerente, Vendedor, Administrador", True),
    (VEN, "OTD-VEN-09",
     "Saber con qué formas de pago cobran las ventas (efectivo, tarjeta, transferencia) "
     "y cómo cambia esa mezcla en el tiempo.",
     False,
     "pago.metodo_pago_id → metodo_pago.nombre/tipo (3 métodos en uso) + pago.monto/"
     "fecha_pago (4 079 pagos, $5 467 791,59 cobrados).",
     FH, "Participación por forma de pago, por mes",
     "Gerente, Analista, Administrador", False),
    (VEN, "OTD-VEN-10",
     "Atender a tiempo la voz del cliente: reseñas en espera de aprobación y preguntas "
     "sobre productos sin responder.",
     True,
     "resena.estado (344 reseñas, 53 pendientes de moderación) + pregunta_producto.estado "
     "(49 preguntas, 13 pendientes) y existencia de respuesta_pregunta (29).",
     FH, "Cola de moderación con antigüedad",
     "Administrador, Gerente (moderadores del sistema)", True),
    (VEN, "OTD-VEN-11",
     "Conocer la calificación que los clientes dan a cada producto y cómo evoluciona.",
     False,
     "resena.calificacion/producto_id/fecha_creacion/compra_verificada — 344 reseñas sobre "
     "268 productos distintos, repartidas en 18 meses (antes 4 reseñas).",
     FH, "Ranking de productos por calificación", "Gerente, Vendedor, Analista", False),
    (VEN, "OTD-VEN-12",
     "Saber cuántos cobros en línea fallan y por qué motivo, para no perder ventas en "
     "el paso del pago.",
     False,
     "Brecha CERRADA: pago.estado registra hoy 'completado' y 'fallido' (176 fallidos) y "
     "transaccion_pago.tipo distingue 'autorizacion' (176, todas de pagos fallidos) de "
     "'captura' (3 903); respuesta_pasarela guarda el motivo — 6 motivos reales: "
     "fondos_insuficientes 40, datos_incorrectos 38, tarjeta_rechazada 37, error_pasarela "
     "31, limite_excedido 28.",
     FH, "Tabla de intentos fallidos por motivo y período", "Gerente, Administrador", False),
    (VEN, "OTD-VEN-13",
     "Saber cuánto vende cada canal — mostrador, teléfono y tienda en línea — y qué "
     "parte de la venta total pone cada uno, por período.",
     False,
     "pedido.canal/total/fecha_pedido — poblados en los 4 083 pedidos: web $3 073 238,46; "
     "mostrador $1 438 538,94; teléfono $1 204 659,15.",
     FH, "Participación de cada canal en la venta, por mes",
     "Gerente, Vendedor, Analista, Administrador", False),
    (VEN, "OTD-VEN-14",
     "Saber cuánto dinero devuelven los clientes al mes y qué porcentaje de la venta "
     "representa, para frenar a tiempo si se dispara.",
     False,
     "devolucion.monto_total/fecha_creacion (196 devoluciones, $95 693,89, mantenido por el "
     "trigger fn_recalcular_total_devolucion) contra pedido.total/fecha_pedido (4 083 "
     "pedidos, $5 716 436,55).",
     FH, "Valor devuelto y porcentaje sobre la venta, mensual",
     "Gerente, Administrador, Analista (Bodega y Despacho NO: es dinero)", False),
    (VEN, "OTD-VEN-15",
     "Seguir la venta acumulada del período contra la meta que se fijó, para reaccionar "
     "a media quincena y no enterarse al cierre del mes.",
     True,
     "Brecha CERRADA: tabla NUEVA meta_venta (anio, mes, departamento, monto_meta, "
     "fijada_por, activo) con 133 filas = 7 departamentos × 19 meses; la venta real se "
     "calcula desde factura_venta no anulada del mes.",
     FH, "Avance de venta contra la meta del período",
     "Gerente, Vendedor, Administrador", True),

    # ── COMPRAS ───────────────────────────────────────────────────────────
    (COM, "OTD-COM-01",
     "Ver las órdenes de compra que esperan aprobación y el estado de cada orden en "
     "curso.",
     True,
     "orden_compra.estado/numero/fecha_emision/total + proveedor.razon_social (865 órdenes "
     "en 19 meses, 11 proveedores). El estado sintético 'pendiente_aprobacion' agrupa "
     "'borrador'+'enviada': aprobar deja la orden en 'confirmada', no existe 'aprobada'.",
     FH, "Tabla con filtros por estado y proveedor",
     "Compras, Gerente, Administrador", True),
    (COM, "OTD-COM-02",
     "Controlar cuánto le debemos a cada proveedor y qué cuotas están por vencer o ya "
     "vencieron.",
     True,
     "cuenta_por_pagar.saldo_pendiente/fecha_vencimiento/estado/proveedor_id (839 cuentas, "
     "276 con saldo abierto, $6 382 924,53) + proveedor.razon_social. Conviven dos "
     "clasificaciones: 'estado' (columna real) y 'situacion' (recalculada hoy contra "
     "fecha_vencimiento).",
     FH, "Tabla de vencimientos con semáforo de fechas",
     "Compras, Gerente, Administrador", True),
    (COM, "OTD-COM-03",
     "Saber si pagamos a los proveedores a tiempo: pagos hechos antes o después de la "
     "fecha de vencimiento, por proveedor y por mes.",
     False,
     "pago_proveedor.fecha_pago/monto/cuenta_por_pagar_id (902 pagos, $16 084 462,74) vs "
     "cuenta_por_pagar.fecha_vencimiento — ambas 100 % pobladas.",
     FH, "Puntualidad de pago por proveedor y mes",
     "Compras, Gerente, Administrador, Analista", False),
    (COM, "OTD-COM-04",
     "Conocer cuánto gastamos en compras por proveedor y por mes.",
     False,
     "factura_compra.total/fecha_emision/proveedor_id (839 facturas, $22 467 387,27, "
     "11 proveedores, 19 meses).",
     FH, "Barras por proveedor con evolución mensual",
     "Compras, Gerente, Administrador, Analista", False),
    (COM, "OTD-COM-05",
     "Saber si cada proveedor cumple el compromiso que pactó: comparar la fecha de "
     "entrega que prometió al confirmar la orden contra el día en que la mercancía "
     "llegó de verdad, para detectar a quién incumple su palabra.",
     False,
     "Brecha CERRADA y ya poblada: orden_compra.fecha_entrega_esperada en 849 de 865 "
     "órdenes, con 825 pares promesa/llegada comparables contra recepcion_mercancia."
     "fecha_recepcion (antes: 0 pares).",
     FH, "Cumplimiento de plazo por proveedor", "Compras, Gerente", False),
    (COM, "OTD-COM-06",
     "Medir el tiempo real observado del ciclo de compra: cuántos días tarda en la "
     "práctica la mercancía en llegar desde que emitimos la orden.",
     False,
     "orden_compra.fecha_emision + recepcion_mercancia.fecha_recepcion/orden_compra_id "
     "(839 recepciones reales).",
     FH, "Días de ciclo de compra por proveedor y período",
     "Compras, Gerente, Analista", False),
    (COM, "OTD-COM-07",
     "Conocer cuánta mercancía llega en mal estado y se rechaza en la puerta al "
     "recibirla, por proveedor y por motivo.",
     False,
     "recepcion_detalle.cantidad_rechazada/motivo_rechazo/cantidad_recibida — 92 de 2 855 "
     "líneas con rechazo registrado y motivo, sobre 6 motivos distintos (antes: 1 de 22).",
     FH, "Porcentaje rechazado por proveedor",
     "Compras, Gerente; Bodega en cantidades, sin montos", False),
    (COM, "OTD-COM-08",
     "Ver los artículos defectuosos pendientes de devolver al proveedor y en qué paso "
     "va cada devolución.",
     True,
     "item_defectuoso.estado/origen/cantidad/proveedor_id (38 ítems, 10 pendientes) + "
     "devolucion_proveedor.numero/estado (8 devoluciones).",
     FH, "Tablero del pool de defectuosos y devoluciones en curso",
     "Compras, Gerente; Bodega en cantidades, sin montos", True),
    (COM, "OTD-COM-09",
     "Saber cuánto recuperamos de los proveedores por mercancía defectuosa: crédito a "
     "favor o reposición de producto.",
     False,
     "devolucion_proveedor.tipo_resolucion/monto_credito/fecha_resolucion + "
     "item_defectuoso.costo_unitario — 8 devoluciones, 6 resueltas (3 reposición + "
     "3 nota de crédito por $4 196,85), repartidas en apenas 6 meses distintos.",
     RV, "Monto recuperado por proveedor y período",
     "Compras, Gerente, Administrador", False),
    (COM, "OTD-COM-10",
     "Comparar a qué proveedor conviene comprarle cada producto: costo, plazo de "
     "entrega y proveedor preferido.",
     True,
     "Brecha CERRADA y poblada: producto_proveedor con 1 106 ofertas de 11 proveedores, "
     "1 093 con tiempo_entrega_dias y 1 040 marcadas es_preferido (antes: 0 filas).",
     FH, "Ficha comparativa de proveedores por producto", "Compras, Gerente", True),
    (COM, "OTD-COM-11",
     "Detectar qué proveedores entregan incompleto: comparar lo que se pidió contra lo "
     "que de verdad llegó, línea por línea y por proveedor.",
     True,
     "orden_compra_detalle.cantidad/cantidad_recibida (259 de 2 949 líneas con recepción "
     "menor a la pedida) + orden_compra.proveedor_id/fecha_emision → proveedor.razon_social.",
     FH, "Líneas incompletas y porcentaje de cumplimiento por proveedor",
     "Compras, Gerente; Bodega en cantidades, sin montos", False),
    (COM, "OTD-COM-12",
     "Saber si está subiendo el costo de lo que compramos: cómo cambia el precio que "
     "cobra el proveedor por cada producto entre una compra y la siguiente.",
     False,
     "orden_compra_detalle.precio_unitario/producto_variante_id (2 949 líneas en 19 meses) "
     "+ orden_compra.fecha_emision/proveedor_id — cada línea de compra conserva su precio "
     "a esa fecha.",
     FH, "Evolución del costo de compra por producto y proveedor",
     "Compras, Gerente, Analista", False),

    # ── INVENTARIO / BODEGA ───────────────────────────────────────────────
    (INV, "OTD-INV-01",
     "Detectar los productos que están por debajo de su tope mínimo y hay que reponer.",
     True,
     "inventario.stock_actual vs inventario.stock_minimo — 1 227 de 1 406 posiciones con "
     "mínimo definido y 162 bajo mínimo hoy. El default 0 significa «sin mínimo definido», "
     "por eso el informe exige stock_minimo > 0.",
     FH, "Lista de reposición con faltante por bodega",
     "Bodega, Compras, Gerente, Administrador", True),
    (INV, "OTD-INV-02",
     "Consultar las existencias actuales de cada producto en cada bodega, incluyendo lo "
     "apartado para pedidos.",
     True,
     "inventario.stock_actual/stock_reservado/bodega_id (1 406 filas) + bodega.nombre "
     "(2 bodegas).",
     FH, "Tabla de existencias con buscador",
     "Bodega, Compras, Vendedor, Gerente, Administrador", True),
    (INV, "OTD-INV-03",
     "Revisar la historia completa de entradas y salidas de cualquier producto: qué se "
     "movió, cuándo, cuánto y por qué razón.",
     True,
     "movimiento_inventario.cantidad/stock_anterior/stock_nuevo/fecha_creacion/"
     "referencia_tipo + tipo_movimiento.codigo/nombre/factor — 13 287 movimientos, los "
     "9 tipos en uso.",
     FH, "Kardex por producto con filtro de tipo y fecha",
     "Bodega, Gerente, Administrador", True),
    (INV, "OTD-INV-04",
     "Saber qué categorías de producto rotan más y cuáles se quedan paradas en bodega, "
     "por período.",
     False,
     "movimiento_inventario.cantidad/fecha_creacion/tipo_movimiento_id (13 287) + "
     "producto_categoria.categoria_id → categoria.nombre (1 214 asignaciones, "
     "11 categorías).",
     FH, "Rotación por categoría y período",
     "Gerente, Analista, Administrador; Bodega en cantidades", False),
    (INV, "OTD-INV-05",
     "Controlar la mercancía perdida o sobrante detectada en los ajustes de inventario "
     "y sus motivos.",
     True,
     "ajuste_inventario.tipo/motivo/estado/fecha_aplicacion — 53 ajustes (50 aplicados, "
     "3 anulados) repartidos en 19 meses + movimientos entrada_ajuste/salida_ajuste del "
     "kardex, enlazados por referencia_tipo.",
     FH, "Lista de ajustes con motivo y cantidades",
     "Bodega, Gerente, Administrador", True),
    (INV, "OTD-INV-06",
     "Seguir las transferencias de mercancía entre bodegas: cuáles van en camino y "
     "cuáles ya se recibieron.",
     True,
     "transferencia_bodega.estado/fecha_envio/fecha_recepcion/bodega_origen_id/"
     "bodega_destino_id — 71 transferencias en los 4 estados (57 recibidas, más en "
     "tránsito, pendientes y canceladas), en 19 meses.",
     FH, "Tabla de transferencias por estado", "Bodega, Gerente, Administrador", True),
    (INV, "OTD-INV-07",
     "Saber cuánto dinero hay parado en mercancía almacenada, por categoría y por "
     "bodega.",
     True,
     "inventario.stock_actual × producto_variante.costo (costo poblado en las 1 221 "
     "variantes, con bandas mayoristas por categoría) + producto_categoria/categoria."
     "nombre — $22 024 063,50 valorizados.",
     FH, "Valor del inventario por categoría/bodega",
     "Gerente, Administrador, Analista (Bodega NO: es dinero)", True),
    (INV, "OTD-INV-08",
     "Detectar productos con demasiada existencia — por encima del tope máximo deseado "
     "— para no enterrar dinero en mercancía de más.",
     True,
     "Brecha CERRADA y poblada: inventario.stock_maximo en 1 227 de 1 406 posiciones, con "
     "184 en sobre-stock hoy (antes: 0 de 1 227, 100 % NULL).",
     FH, "Lista de sobre-stock por bodega", "Bodega, Compras, Gerente", True),
    (INV, "OTD-INV-09",
     "Ver cómo evoluciona mes a mes el dinero inmovilizado en la mercancía almacenada, "
     "para saber si la bodega se está llenando o vaciando de capital.",
     False,
     "Reconstrucción del stock al cierre de cada mes: inventario.stock_actual (1 406 filas) "
     "menos los movimientos posteriores del kardex — movimiento_inventario.cantidad/"
     "fecha_creacion (13 287, encadenados por (fecha_creacion, id) con stock_anterior/"
     "stock_nuevo como respaldo) y el signo de tipo_movimiento.factor — valorizado con "
     "producto_variante.costo. Usa el costo vigente: no existe histórico de costos.",
     FH, "Línea mensual del valor almacenado por bodega y categoría",
     "Gerente, Administrador, Analista (Bodega NO: es dinero)", False),
    (INV, "OTD-INV-10",
     "Conocer las mermas (mercancía perdida) y los sobrantes acumulados por período y "
     "por motivo, para atacar las causas de la pérdida.",
     False,
     "ajuste_inventario.tipo/motivo/fecha_aplicacion (53 ajustes sobre 7 motivos "
     "tipificados, en 19 meses) + kardex movimiento_inventario.cantidad con "
     "tipo_movimiento.codigo 'salida_ajuste'/'entrada_ajuste', enlazados por "
     "referencia_tipo='ajuste_inventario'.",
     FH, "Acumulado por motivo y mes",
     "Bodega en cantidades; valorizado solo Gerente, Administrador", False),

    # ── LOGÍSTICA / DESPACHO ──────────────────────────────────────────────
    (LOG, "OTD-LOG-01",
     "Ver la cola de pedidos listos y en espera de ser despachados.",
     True,
     "pedido.estado_pedido_id → estado_pedido.codigo de los tres estados del tramo de "
     "salida (facturado / en_preparacion / preparado) + pedido.numero/fecha_pedido — "
     "48 pedidos en cola hoy.",
     FH, "Cola de trabajo de despacho", "Despacho, Gerente, Administrador", True),
    (LOG, "OTD-LOG-02",
     "Seguir los envíos: cuáles van en camino, cuáles ya se entregaron, cuáles "
     "volvieron, y con qué transportista y número de guía viaja cada uno.",
     True,
     "envio.estado/numero_guia/fecha_despacho (2 872 envíos, 2 727 con entrega real) + "
     "transportista.nombre (5 transportistas).",
     FH, "Tablero de envíos por estado y transportista",
     "Despacho, Gerente, Administrador", True),
    (LOG, "OTD-LOG-03",
     "Saber si cumplimos la fecha de entrega prometida al cliente: de los envíos ya "
     "entregados, cuántos llegaron a más tardar el día prometido, por transportista.",
     False,
     "envio.fecha_entrega_estimada vs envio.fecha_entrega_real — 2 723 envíos con ambas "
     "fechas + envio.transportista_id (5 transportistas).",
     FH, "Cumplimiento de fecha prometida por transportista",
     "Despacho, Gerente, Analista", False),
    (LOG, "OTD-LOG-04",
     "Medir la duración real del tránsito — los días desde que el paquete sale de "
     "bodega hasta la puerta del cliente — para comparar transportistas entre sí.",
     False,
     "envio.fecha_despacho vs envio.fecha_entrega_real (ambas pobladas en los 2 727 "
     "envíos entregados).",
     FH, "Días de tránsito promedio por transportista y período",
     "Despacho, Gerente, Analista", False),
    (LOG, "OTD-LOG-05",
     "Controlar los problemas de entrega — cliente ausente, dirección equivocada, "
     "rechazo en la puerta, daño en el camino — cuántos ocurren, cuántos intentos "
     "toman y cómo terminan.",
     False,
     "novedad_envio.tipo/estado/accion/fecha_registro/fecha_resolucion — 176 novedades "
     "(169 resueltas) repartidas en 19 meses sobre los 5 tipos del CHECK; los intentos se "
     "calculan como 1 + reprogramaciones.",
     FH, "Problemas de entrega por tipo y desenlace",
     "Despacho, Gerente, Soporte", False),
    (LOG, "OTD-LOG-06",
     "Ver las devoluciones de clientes en curso y en qué paso va cada una (solicitada, "
     "en revisión, aprobada, en camino de vuelta, recibida, inspeccionada, "
     "reembolsada, cerrada).",
     True,
     "devolucion.numero/estado/guia_retorno/fecha_creacion (196 devoluciones, 91 en "
     "estados terminales) + devolucion_detalle.resultado_inspeccion (162 líneas "
     "inspeccionadas).",
     FH, "Tablero del ciclo de devolución por estado",
     "Despacho, Soporte, Gerente; Bodega ve su tramo de inspección, sin montos", True),
    (LOG, "OTD-LOG-07",
     "Medir cuántos días tarda una devolución de cliente desde que se solicita hasta "
     "que se cierra.",
     False,
     "devolucion.fecha_creacion + historial_estado_devolucion.estado/fecha_creacion — "
     "1 008 registros de historial; 161 de las 196 devoluciones tienen 3 o más pasos "
     "fechados, así que los tramos intermedios son medibles.",
     FH, "Días de ciclo de devolución por período",
     "Gerente, Soporte, Analista", False),
    (LOG, "OTD-LOG-08",
     "Saber por qué devuelven los clientes y qué pasa con esa mercancía: vuelve a "
     "venderse, resulta defectuosa o se rechaza sin reembolso.",
     False,
     "devolucion.motivo_devolucion_id → motivo_devolucion.nombre (los 4 motivos en uso) + "
     "devolucion_detalle.resultado_inspeccion/cantidad (274 líneas, 162 con inspección "
     "registrada).",
     FH, "Motivos de devolución y destino de la mercancía",
     "Gerente, Soporte, Analista; Bodega en cantidades", False),
    (LOG, "OTD-LOG-09",
     "Saber, de cada 100 envíos, cuántos terminan en devolución, mes a mes.",
     False,
     "envio.fecha_despacho (2 872 envíos) vs devolucion.pedido_id/fecha_creacion "
     "(196 devoluciones), ambos con cobertura en los 19 meses.",
     FH, "Proporción de devoluciones sobre envíos, mensual",
     "Gerente, Analista; Despacho en conteos", False),
    (LOG, "OTD-LOG-10",
     "Controlar el dinero devuelto a los clientes: cuánto se reembolsó, por qué vía y "
     "por qué motivo.",
     False,
     "Brecha CERRADA y poblada: tabla reembolso con 85 filas y $44 525,63 en monto, vía y "
     "fecha, enlazadas a su devolución (antes: 1 fila).",
     FH, "Reembolsos pagados por período y vía",
     "Gerente, Administrador, Soporte (Despacho NO: es dinero)", False),
    (LOG, "OTD-LOG-11",
     "Saber cuánto nos cuesta llevar cada envío, por zona y transportista, para "
     "revisar las tarifas que cobramos al cliente.",
     True,
     "Brecha CERRADA y poblada: envio.costo > 0 en 2 848 de 2 872 envíos ($32 723,25 en "
     "total), con peso_total_kg en 2 848; producto_variante.peso_kg poblado en las 1 221 "
     "variantes y tarifa_envio.costo_por_kg > 0 en las 3 tarifas. La ZONA no es una "
     "columna del envío: se resuelve desde la dirección del pedido por ciudad > provincia "
     "> país, la misma cadena que asigna el transportista.",
     FH, "Costo de envío por zona y transportista",
     "Gerente, Administrador (Despacho NO: es dinero)", True),
    (LOG, "OTD-LOG-12",
     "Medir cuánto tarda un pedido en cada etapa del camino — del pago a la "
     "preparación, de la preparación al despacho y del despacho a la entrega — para "
     "encontrar el cuello de botella real y no suponerlo.",
     False,
     "historial_estado_pedido.pedido_id/estado_pedido_id/fecha_creacion (24 608 registros "
     "que cubren los 4 083 pedidos) + estado_pedido.codigo — cada transición quedó fechada, "
     "lo que permite medir cada tramo.",
     FH, "Días u horas promedio por etapa del ciclo, por período",
     "Despacho (tiempos y estados, sin dinero), Gerente, Analista", False),

    # ── SOPORTE ───────────────────────────────────────────────────────────
    (SOP, "OTD-SOP-01",
     "Ver los reclamos y consultas abiertos: cuántos hay, de qué tipo, qué tan "
     "urgentes son y quién atiende cada uno.",
     True,
     "ticket_soporte.numero/estado/prioridad/categoria_ticket_id/asignado_usuario_id — "
     "248 tickets, 128 vivos, 6 agentes con casos asignados.",
     FH, "Bandeja con filtros por estado, urgencia y categoría",
     "Soporte, Gerente, Administrador", True),
    (SOP, "OTD-SOP-02",
     "Saber cuántos reclamos atendemos dentro del tiempo prometido al cliente según su "
     "urgencia.",
     False,
     "Brecha CERRADA: ticket_soporte.fecha_limite existe y está poblada en 248 de 248 "
     "tickets (con backfill derivado del plazo por urgencia 2 h/4 h/24 h/72 h, que antes "
     "vivía solo en código) y se contrasta con fecha_cierre (76 tickets cerrados).",
     FH, "Cumplimiento del tiempo prometido por período", "Soporte, Gerente", False),
    (SOP, "OTD-SOP-03",
     "Medir cuántas horas o días tardamos en resolver cada tipo de problema.",
     False,
     "ticket_soporte.fecha_creacion vs ticket_soporte.fecha_cierre (76 tickets cerrados, "
     "antes 2) + categoria_ticket.nombre (8 categorías, todas con casos).",
     FH, "Tiempo de resolución por categoría", "Soporte, Gerente, Analista", False),
    (SOP, "OTD-SOP-04",
     "Saber qué tipos de problema generan más reclamos, para atacar la causa de fondo.",
     True,
     "ticket_soporte.categoria_ticket_id → categoria_ticket.nombre — las 8 categorías "
     "tienen tickets (antes 4 de 8 estaban vacías); el porcentaje se calcula sobre el "
     "total del período.",
     FH, "Distribución de tickets por categoría", "Soporte, Gerente", True),
    (SOP, "OTD-SOP-05",
     "Repartir bien el trabajo del equipo: reclamos asignados y resueltos por cada "
     "persona de soporte.",
     True,
     "ticket_soporte.asignado_usuario_id → usuario.nombre/apellido (6 agentes) + "
     "ticket_soporte.estado; la fila «(sin asignar)» se emite aparte porque es el dato "
     "accionable.",
     FH, "Carga y cierre por agente", "Soporte, Gerente, Administrador", True),
    (SOP, "OTD-SOP-06",
     "Medir cuánto tardamos en dar la primera respuesta al cliente que reclama: desde "
     "que abre su caso hasta que alguien del equipo le contesta por primera vez.",
     False,
     "ticket_soporte.fecha_creacion vs el primer mensaje_ticket.fecha_creacion cuyo autor "
     "es del equipo (usuario_id poblado) y visible al cliente (es_interno = false) — 193 de "
     "248 tickets con primera respuesta del equipo, sobre 523 mensajes.",
     FH, "Horas hasta la primera respuesta, por período y urgencia",
     "Soporte, Gerente", False),
    (SOP, "OTD-SOP-07",
     "Medir cuánto tarda cada persona del equipo de soporte en resolver los casos que "
     "se le asignan.",
     False,
     "ticket_soporte.asignado_usuario_id → usuario.nombre/apellido (6 agentes) + "
     "ticket_soporte.fecha_creacion/fecha_cierre (76 cierres reales, antes 2).",
     FH, "Tiempo de resolución por agente y período", "Soporte, Gerente", False),
    (SOP, "OTD-SOP-08",
     "Saber qué productos generan más reclamos y más devoluciones, para pedir a "
     "Compras que revise el producto o a Ventas que corrija su descripción.",
     False,
     "Brecha CERRADA: ticket_soporte.producto_variante_id existe y está poblada en 142 de "
     "248 tickets (antes la columna no existía). La otra mitad ya era factible: "
     "devoluciones por producto vía devolucion_detalle.pedido_detalle_id → pedido_detalle."
     "producto_variante_id → producto.nombre (274 líneas).",
     FH, "Ranking de productos por reclamos y devoluciones",
     "Soporte, Gerente, Compras", False),

    # ── GERENCIA / DIRECCIÓN ──────────────────────────────────────────────
    (GER, "OTD-GER-01",
     "Tener la foto del día del negocio: pedidos de hoy, dinero cobrado hoy y "
     "pendientes que necesitan decisión.",
     True,
     "pedido.fecha_pedido/total/estado_pedido_id (4 083) + pago.fecha_pago/monto (4 079) + "
     "factura_venta (3 887) + estado_pedido.nombre. El bloque de pendientes es AL MOMENTO, "
     "no del día consultado.",
     FH, "Tarjetas de resumen del día", "Gerente, Administrador", True),
    (GER, "OTD-GER-02",
     "Comparar mes a mes el dinero que entra por ventas contra el dinero que sale "
     "hacia proveedores — la balanza comercial interna del negocio.",
     False,
     "Facturado: factura_venta.total/fecha_emision (3 887 facturas, $5 417 807,65 sin "
     "anuladas) vs factura_compra.total/fecha_emision (839, $22 467 387,27). En caja: "
     "pago.monto/fecha_pago ($5 467 791,59) vs pago_proveedor.monto/fecha_pago "
     "($16 084 462,74).",
     FH, "Barras enfrentadas entradas vs salidas, por mes",
     "Gerente, Administrador, Analista", False),
    (GER, "OTD-GER-03",
     "Saber qué categorías de producto dejan más ganancia (venta menos costo), por "
     "período.",
     False,
     "pedido_detalle.precio_unitario/cantidad/monto_descuento (10 384 líneas) + "
     "producto_variante.costo (1 221 variantes, bandas mayoristas por categoría) + "
     "producto_categoria/categoria.nombre (11 categorías).",
     FH, "Ganancia por categoría y período", "Gerente, Analista, Administrador", False),
    (GER, "OTD-GER-04",
     "Ver los cupones de descuento activos, cuántos usos les quedan y cuándo vencen.",
     True,
     "cupon.codigo/activo/usos_actuales/usos_maximos/fecha_inicio/fecha_fin — 33 cupones, "
     "7 vigentes. La 'situacion' replica las TRES condiciones de canje del motor de "
     "descuentos (activo + ventana + usos disponibles), no solo 'activo'.",
     FH, "Tabla de cupones con vigencia y usos restantes",
     "Gerente, Administrador", True),
    (GER, "OTD-GER-05",
     "Saber qué cupones usaron efectivamente los clientes y cuánto descuento le "
     "costaron al negocio, por período.",
     False,
     "uso_cupon.cupon_id/monto_descontado/fecha_creacion — 564 usos de 25 cupones "
     "distintos, $50 727,89 otorgados, repartidos en 19 meses (antes: 3 usos, $52,91).",
     FH, "Descuento otorgado por cupón y período",
     "Gerente, Analista, Administrador", False),
    (GER, "OTD-GER-06",
     "Ver las acciones de marketing vigentes: promociones (y los productos que "
     "abarcan), campañas y anuncios de la tienda.",
     True,
     "promocion.fecha_inicio/fecha_fin/activo (24 promociones, 6 vigentes) + "
     "promocion_producto.producto_id (232 asignaciones, las 24 promos con productos) + "
     "campana.estado/fecha_inicio/fecha_fin (18, 6 activas) + banner.activo/fecha_inicio "
     "(23, 8 activos) — 20 elementos vigentes de 65.",
     FH, "Panel de vigencias de marketing", "Gerente, Administrador", True),
    (GER, "OTD-GER-07",
     "Saber si las promociones hacen vender más, comparando las ventas del producto "
     "antes y durante la promoción.",
     False,
     "pedido_detalle.monto_descuento (>0 = línea con promoción aplicada) + promocion."
     "fecha_inicio/fecha_fin + promocion_producto — hay 24 promociones con 232 productos y "
     "línea base amplia (4 133 líneas de esos productos ANTES de su ventana), pero solo "
     "123 líneas efectivamente promocionadas: la muestra «durante» es demasiado pequeña "
     "frente a la base.",
     RV, "Ventas antes/durante promoción", "Gerente, Analista", False),
    (GER, "OTD-GER-08",
     "Revisar quién hizo qué en el sistema: aprobaciones, despachos, registros y "
     "moderaciones, con autor, fecha y detalle del cambio.",
     True,
     "log_auditoria.usuario_id/tabla/accion/registro_id/datos_anteriores/datos_nuevos/"
     "fecha_creacion — 7 073 registros sobre 10 tablas, en 19 meses. DATO SENSIBLE: solo "
     "Administrador y Gerente.",
     FH, "Registro de acciones filtrable por usuario, tabla y fecha",
     "Administrador, Gerente", True),
    (GER, "OTD-GER-09",
     "Saber quién intentó entrar al sistema y falló, desde dónde y por qué.",
     True,
     "Brecha CERRADA y poblada: log_acceso.exitoso/motivo_fallo/ip_origen/email_intentado/"
     "fecha_creacion con 1 537 intentos, 201 fallidos sobre los 4 motivos de rechazo del "
     "login, en 19 meses (antes: 0 filas). DATO SENSIBLE: solo Administrador y Gerente.",
     FH, "Intentos de acceso fallidos recientes", "Administrador, Gerente", True),
    (GER, "OTD-GER-10",
     "Conocer la ganancia real de cada producto — la diferencia entre lo que costó y "
     "lo que se cobró — además de la vista por categoría, para corregir precios "
     "producto por producto.",
     False,
     "producto_variante.precio/costo (ambos poblados en las 1 221 variantes) + "
     "pedido_detalle.producto_variante_id/cantidad/precio_unitario/monto_descuento "
     "(10 384 líneas, 834 variantes con venta). Usa el costo vigente: no existe histórico "
     "de costos.",
     FH, "Margen por producto con buscador y filtro de período",
     "Gerente, Analista, Administrador", False),
    (GER, "OTD-GER-11",
     "Controlar cuánto descuento se entrega en total cada mes — sumando promociones y "
     "cupones — y sobre qué productos recae, para que el descuento no se coma el "
     "margen sin que nadie lo mida.",
     False,
     "pedido_detalle.monto_descuento (123 líneas con promoción) + pedido.monto_descuento "
     "(562 pedidos con cupón) + uso_cupon.monto_descontado (564 usos, $50 727,89), con el "
     "prorrateo por producto ya resuelto en factura_venta_detalle.monto_descuento.",
     FH, "Descuento total entregado por mes y por producto",
     "Gerente, Analista, Administrador", False),
]

# ══════════════ CUADRE ARITMÉTICO (falla ruidosamente si algo no cuadra) ══
assert len(OBJ) == 68, "El catálogo debe tener exactamente 68 objetivos"
N_SIMPLES = sum(1 for o in OBJ if o[3])
N_COMPUESTOS = sum(1 for o in OBJ if not o[3])
N_FH = sum(1 for o in OBJ if o[5] == FH)
N_RC = sum(1 for o in OBJ if o[5] == RC)
N_RV = sum(1 for o in OBJ if o[5] == RV)
N_IMPL = sum(1 for o in OBJ if o[8])
assert (N_SIMPLES, N_COMPUESTOS) == (29, 39), "Cuadre simples/compuestos roto"
assert (N_FH, N_RC, N_RV) == (66, 0, 2), "Cuadre de factibilidad roto"
assert N_FH + N_RC + N_RV == 68
assert N_IMPL == 28, "Deben estar marcados 28 informes implementados"
assert all(o[3] for o in OBJ if o[8]), "Solo un informe SIMPLE puede estar implementado"

DEPTOS = [VEN, COM, INV, LOG, SOP, GER]
def _cuenta(d):
    filas = [o for o in OBJ if o[0] == d]
    return (len(filas), sum(1 for o in filas if o[3]), sum(1 for o in filas if not o[3]),
            sum(1 for o in filas if o[8]))
RESUMEN_DEPTO = {d: _cuenta(d) for d in DEPTOS}
assert sum(v[0] for v in RESUMEN_DEPTO.values()) == 68
assert sum(v[1] for v in RESUMEN_DEPTO.values()) == 29
assert sum(v[3] for v in RESUMEN_DEPTO.values()) == 28

# ══════════════════════════════ PORTADA ═══════════════════════════════════
p()
p("UNIVERSIDAD TÉCNICA ESTATAL DE QUEVEDO", bold=True, size=18, align=CENTER)
p("_______________________________________", align=CENTER)
p()
p("Facultad: ______________________________", size=12, align=CENTER)
p("Carrera: _______________________________", size=12, align=CENTER)
p()
p("Asignatura", bold=True, size=12, align=CENTER)
p("Construcción de Software", size=12, align=CENTER)
p()
p("TAREA 11", bold=True, size=16, align=CENTER)
p("Análisis del Nivel Táctico del Sistema RetailMind:", bold=True, size=14, align=CENTER)
p("Objetivos Tácticos por Departamento, Clasificación de Informes", bold=True, size=14, align=CENTER)
p("e Implementación de los Reportes Simples", bold=True, size=14, align=CENTER)
p("(Simples — BDR PostgreSQL / Compuestos — BD Columnar ClickHouse)", size=12, align=CENTER)
p("Versión 3 — catálogo sincronizado y verificado contra la base de datos real",
  size=11, align=CENTER)
p()
p("Estudiante", bold=True, size=12, align=CENTER)
p("Benites Pérez Dariem Alberto", size=12, align=CENTER)
p()
p("Docente: ______________________________", size=12, align=CENTER)
p("Semestre: Sexto — Fecha: ______________", size=12, align=CENTER)
doc.add_page_break()

# ══════════════════════════════ 1. INTRODUCCIÓN ═══════════════════════════
h(1, "1. Introducción")
p("El sistema RetailMind concluyó la construcción de su nivel operativo: el registro de las "
  "transacciones diarias del negocio (ciclo completo de venta con pagos y facturación, ciclo de "
  "compra con aprobaciones, inventario con kardex, logística de despacho y devoluciones, soporte "
  "al cliente y marketing). El presente documento da el siguiente paso en la pirámide "
  "informacional de la organización: el análisis del nivel táctico, y la implementación de los "
  "reportes simples que de él se derivan.")
p("Mientras el nivel operativo responde a la pregunta «¿qué transacción ocurre ahora?», el nivel "
  "táctico responde a «¿cómo se dirige y controla cada área?». Un objetivo táctico es una "
  "necesidad de información planteada por el jefe de cada departamento, orientada a la dirección, "
  "el seguimiento y el control de su área en el mediano plazo. A partir de esos objetivos se "
  "derivan los informes tácticos que el sistema debe producir.")
p("Los objetivos tácticos se levantaron para los SEIS departamentos del negocio: Ventas, "
  "Compras, Inventario/Bodega, Logística/Despacho, Soporte y Gerencia/Dirección (que dirige "
  "también Marketing). El número de objetivos por departamento no es una cuota: sale de lo que "
  "cada área realmente necesita dirigir, y por eso es asimétrico.")

# ── ADVERTENCIA (ver cabecera del archivo) ────────────────────────────────
# ESTA es la seccion refutada. El texto de abajo caracteriza el negocio como
# "distribuidora mayorista B2B"; el diagnostico del 2026-07-30
# (docs/estrategico/DIAGNOSTICO_SEGMENTO_CLIENTE.md, veredicto (c) poblacion
# homogenea) probo que RetailMind es un COMERCIO MINORISTA MULTICANAL DE
# TICKET ALTO ($1.400,06 por pedido formado por un precio unitario de $276,36,
# no por cantidad). Se deja intacto porque el entregable ya se presento; si se
# regenera el documento, este bloque DEBE reescribirse primero.
h(2, "1.1. El negocio que se dirige: distribuidora mayorista B2B")
p("RetailMind opera como distribuidora mayorista ecuatoriana con sede en Quevedo: compra volumen "
  "a proveedores, almacena en dos bodegas y revende a clientes corporativos y minoristas "
  "recurrentes por tres canales (mostrador, teléfono y tienda en línea). Esa caracterización no "
  "es decorativa — condiciona qué informe táctico tiene sentido pedir:")
for txt in [
    "El abastecimiento pesa más que la venta. El histórico verificado registra $22 467 387,27 "
    "facturados en compras contra $5 716 436,55 vendidos, y $6 382 924,53 de saldo abierto en "
    "276 cuentas por pagar vivas. Por eso Compras tiene 12 objetivos y la mitad miran deuda, "
    "puntualidad y cumplimiento del proveedor: en una distribuidora, el margen se gana comprando.",
    "El margen es de distribución, no de detalle. El costo por variante se sitúa en bandas "
    "mayoristas por categoría, de modo que los objetivos de rentabilidad miden puntos de margen "
    "sobre volumen, no ticket unitario.",
    "El cliente es recurrente y de volumen, no un comprador ocasional: 69 de 72 clientes tienen "
    "pedidos y el ticket se forma sobre 10 384 líneas en 4 083 pedidos. Eso vuelve significativos "
    "los objetivos de comportamiento por cliente y de descuento negociado, que en un minorista de "
    "paso serían ruido.",
    "El capital vive en la bodega: $22 024 063,50 de inventario valorizado, varias veces la venta "
    "anual. De ahí que Inventario concentre siete de sus diez objetivos en el control del presente "
    "y que el sobre-stock sea una pregunta de dirección, no un detalle operativo.",
]:
    doc.add_paragraph(txt, style="List Bullet")

h(2, "1.2. Informes simples y compuestos: regla de clasificación")
p("Cada informe táctico se clasifica con una regla explícita, que resuelve los casos ambiguos "
  "(por ejemplo, «ventas por vendedor» y «ventas por cajero», consultas casi idénticas que en el "
  "enunciado aparecían en lados distintos):")
par = doc.add_paragraph(style="List Bullet")
r = par.add_run("Informe SIMPLE: ")
r.bold = True
par.add_run("responde sobre el ESTADO ACTUAL del área. Puede contar o sumar sobre la foto de "
            "hoy, pero NO recorre histórico ni compara períodos. Se resuelve con una consulta "
            "directa a la base de datos relacional (BDR PostgreSQL). Ejemplo: pedidos "
            "preparados en espera de despacho.")
par = doc.add_paragraph(style="List Bullet")
r = par.add_run("Informe COMPUESTO: ")
r.bold = True
par.add_run("requiere recorrer datos históricos, comparar entre períodos o cruzar varias "
            "fuentes con agregación. Se procesa en la base de datos columnar (BDColumnar "
            "ClickHouse), alimentada mediante un proceso ETL orquestado con Apache Airflow. "
            "Ejemplo: tendencia de ventas por mes y categoría.")
p("Dos precisiones para los casos limítrofes: (1) sumar o contar sobre el presente no convierte "
  "un informe en compuesto — «ventas por vendedor» es simple mientras totalice la cartera "
  "actual, y solo se vuelve compuesto cuando compara mes contra mes; (2) una consulta puntual de "
  "detalle que lista filas antiguas sin agregarlas (el historial de movimientos de UN producto, "
  "el registro de auditoría filtrado) sigue siendo simple: es una consulta directa filtrada, no "
  "un barrido agregado del histórico.")
p("Aplicada esta regla, los ocho objetivos simples que contienen agregación (OTD-VEN-02, "
  "OTD-VEN-15, OTD-COM-11, OTD-INV-07, OTD-LOG-11, OTD-SOP-04, OTD-SOP-05 y OTD-GER-01) agregan "
  "sobre la foto presente sin comparar períodos y permanecen SIMPLES. Un solo objetivo cambió de "
  "clasificación en toda la vida del catálogo: OTD-LOG-11 (costo de envío por zona y "
  "transportista) pasó de COMPUESTO a SIMPLE al comprobarse, en su implementación, que la "
  "pregunta que el jefe de despacho realmente hace se responde con una foto agregada del "
  "presente. La serie temporal del costo de envío — cómo evoluciona la tarifa mes a mes, que sí "
  "es compuesta — queda para ClickHouse.")

h(2, "1.3. Arquitectura híbrida del sistema real")
p("RetailMind ya materializa esta separación con una arquitectura híbrida de dos motores: "
  "PostgreSQL como única base transaccional (BDR, base retailmind con 110 tablas operativas en "
  "el esquema public — conteo verificado contra el catálogo del motor —, seguridad por roles, "
  "RLS y triggers de negocio) y ClickHouse como base columnar dedicada exclusivamente a "
  "analítica. Los informes simples de este análisis se resuelven con consultas directas sobre "
  "PostgreSQL, y 28 de ellos ya están construidos y consultables en la aplicación (sección 6); "
  "los compuestos justifican el pipeline ETL (extracción desde la BDR, transformación y carga en "
  "ClickHouse) cuya orquestación con Airflow constituye la fase pendiente del proyecto.")

# ══════════════════════════════ 2. METODOLOGÍA ════════════════════════════
h(1, "2. Metodología")
p("El análisis siguió el proceso enseñado en clase, en este orden estricto:")
par = doc.add_paragraph(style="List Number")
par.add_run("Primero se levantaron las preguntas de negocio de cada jefatura SIN mirar el "
            "esquema de la base de datos: cada jefe departamental, como proveedor de objetivos "
            "tácticos, dijo qué necesita para dirigir y controlar su área, en su propio "
            "lenguaje. En esta fase se levantaron 78 preguntas de negocio a ciegas, cada una "
            "acompañada de la decisión concreta que su respuesta permite tomar. El cruce de esas "
            "78 preguntas contra el catálogo detectó 20 puntos ciegos, de los cuales se "
            "incorporaron los 11 de mayor valor de dirección.")
par = doc.add_paragraph(style="List Number")
par.add_run("Después se verificó contra la base de datos real cuáles de esas preguntas puede "
            "responder: tabla por tabla y columna por columna se confirmó que el dato exista Y "
            "esté poblado — no basta que la columna esté declarada en el esquema si ningún "
            "flujo la escribe. Toda cita de tablas y columnas de este documento proviene de "
            "consultas reales de solo lectura ejecutadas contra la base retailmind.")
par = doc.add_paragraph(style="List Number")
par.add_run("Donde la base dijo que NO, se corrigió en TRES capas antes de construir el ETL: la "
            "base de datos (columna o tabla), el backend (el flujo que debe escribir el dato) y "
            "el formulario de la aplicación (donde el usuario lo captura). No basta con agregar "
            "la columna: si nadie la llena desde la pantalla, el informe seguirá vacío. Las diez "
            "brechas detectadas a lo largo del análisis se cerraron así, una por una (sección 5).")
par = doc.add_paragraph(style="List Number")
par.add_run("Con el dato disponible, se construyeron los informes SIMPLES sobre la BDR y se "
            "pusieron a consultar por pantalla (sección 6). Los COMPUESTOS quedan para la fase "
            "ETL hacia ClickHouse.")
p("El resultado son 68 objetivos tácticos verificados: 66 respondibles hoy, 0 que exijan un "
  "cambio en el sistema — las diez brechas históricas están cerradas — y 2 cuyo esquema y flujo "
  "están completos pero cuya densidad de operación aún no sostiene un agregado representativo.")

p()
par = doc.add_paragraph()
r = par.add_run("Evolución del estado de factibilidad a lo largo del análisis: ")
r.bold = True
r.italic = True
table(
    ["CORTE", "FACTIBLE HOY", "REQUIERE CAMBIO EN EL SISTEMA", "REQUIERE VOLUMEN DE DATOS"],
    [
        ("Catálogo ampliado inicial", "44", "10", "14"),
        ("Tras el cierre de las 4 primeras brechas", "45", "6", "17"),
        ("Versión actual (verificada)", "66", "0", "2"),
    ],
    widths=[2.3, 1.3, 1.9, 1.6],
    font_size=9,
)
p()
p("El salto de 45 a 66 tiene dos causas verificables: (a) el cierre de las seis brechas de "
  "sistema restantes —metas de venta, registro de cobros fallidos, catálogo proveedor-producto, "
  "fecha límite del ticket, producto asociado al reclamo y registro de intentos de acceso—, cada "
  "una en sus tres capas; y (b) la carga de 19 meses de operación histórica (4 083 pedidos, 865 "
  "órdenes de compra, 2 872 envíos, 248 tickets, 344 reseñas, 13 287 movimientos de kardex, "
  "7 073 registros de auditoría y 1 537 intentos de acceso), que llenó los volúmenes que "
  "faltaban para que los informes agregados fueran representativos.")

# ═══════════ 3. TABLA PRINCIPAL DE OBJETIVOS TÁCTICOS ═════════════════════
doc.add_page_break()
h(1, "3. Objetivos tácticos por departamento y clasificación de informes")
p("La siguiente tabla consolida los 68 objetivos tácticos provistos por las seis jefaturas "
  "departamentales, la clasificación de su informe (simple o compuesto) y la base de datos de "
  "la que se obtiene cada uno. Es el cuerpo central del entregable.")

filas_principal = [
    (o[0], o[2], "X" if o[3] else "—", "—" if o[3] else "X", BDR if o[3] else BDC)
    for o in OBJ
]
t = table(
    ["DEPARTAMENTO", "OBJETIVO TÁCTICO", "¿ES UN INFORME SIMPLE?",
     "¿ES UN INFORME COMPUESTO?", "FUENTE (BDR PostgreSQL / BDColumnar ClickHouse)"],
    filas_principal,
    widths=[1.0, 2.6, 0.7, 0.8, 1.4],
    font_size=8,
)
merge_first_col(t, filas_principal)

# ═══════════ 4. VERIFICACIÓN DE FACTIBILIDAD ══════════════════════════════
doc.add_page_break()
h(1, "4. Verificación de factibilidad sobre la base de datos real")
p("Cada objetivo fue contrastado contra la base retailmind con consultas de solo lectura: se "
  "citan las tablas y columnas concretas que lo sustentan y el conteo real observado, y se "
  "declara su estado. Los estados posibles son: FACTIBLE HOY (el dato existe y está poblado), "
  "REQUIERE CAMBIO EN EL SISTEMA (falta que alguna capa escriba el dato; ver sección 5) y "
  "REQUIERE VOLUMEN DE DATOS (esquema y flujo completos, pero la operación registrada aún no "
  "sostiene un agregado representativo).")

filas_verif = [(o[0], o[1], o[4], o[5]) for o in OBJ]
t = table(
    ["DEPTO.", "ID", "SOPORTE DE DATOS VERIFICADO (tablas.columnas y conteos reales)", "ESTADO"],
    filas_verif,
    widths=[0.8, 1.0, 3.5, 1.2],
    font_size=8,
)
merge_first_col(t, filas_verif, font_size=8)

p()
par = doc.add_paragraph()
r = par.add_run("Nota de verificación: ")
r.bold = True
r.italic = True
r2 = par.add_run(
    "ninguna cita de esta tabla es genérica ni inferida. Cada tabla.columna y cada conteo "
    "(4 083 pedidos, 24 608 transiciones de estado, 13 287 movimientos de kardex, 825 pares "
    "promesa/llegada de proveedor, 2 848 envíos con costo calculado, 193 de 248 tickets con "
    "primera respuesta del equipo, 1 537 intentos de acceso registrados) proviene de una consulta "
    "de solo lectura ejecutada contra la base el 2026-07-26. El criterio se mantuvo también en "
    "sentido contrario: los dos objetivos que la operación todavía no densifica lo suficiente "
    "(OTD-COM-09, con 6 resoluciones de devolución a proveedor sobre 11 proveedores y 19 meses, y "
    "OTD-GER-07, con 123 líneas promocionadas frente a 4 133 de línea base) se declaran "
    "honestamente como REQUIERE VOLUMEN DE DATOS en vez de darlos por cubiertos porque la consulta "
    "«devuelve filas».")
r2.italic = True

# ═══════════ 5. CIERRE DE BRECHAS ═════════════════════════════════════════
doc.add_page_break()
h(1, "5. Cierre de brechas: qué faltaba y en qué capa se corrigió")
p("Esta sección aplica el método completo que el profesor enfatizó: no basta con la base de "
  "datos. Si falta un dato, hay que agregarlo también al flujo que lo escribe y al formulario "
  "donde se captura. A lo largo del análisis, diez objetivos no pudieron responderse porque "
  "alguna de las tres capas no producía el dato. LOS DIEZ ESTÁN HOY CERRADOS, verificados contra "
  "el sistema real y devolviendo filas.")

brechas = [
    ("OTD-VEN-12",
     "Los intentos de cobro fallidos nunca se registraban: pago.estado valía 'completado' en el "
     "100 % de las filas.",
     "Ninguno — pago.estado y transaccion_pago.tipo/respuesta_pasarela ya existían.",
     "La simulación de pasarela del checkout registra el intento rechazado en transacción propia "
     "(REQUIRES_NEW) con su motivo, antes de devolver el error.",
     "El formulario de pago ya capturaba todo; ahora muestra el motivo del rechazo al cliente.",
     "176 pagos fallidos con 6 motivos distintos."),
    ("OTD-VEN-15",
     "No existía ninguna tabla de metas de venta: la venta real estaba, pero no había meta contra "
     "la cual compararla.",
     "Tabla NUEVA meta_venta (anio, mes, departamento, monto_meta, fijada_por, activo), con sus "
     "permisos, RLS y política de horario según el patrón vigente. Es la única tabla nueva que "
     "pidió todo el catálogo.",
     "Servicio de metas: crear, editar y consultar el avance contra la venta real del período.",
     "Formulario de captura de metas para Gerencia (período, departamento, monto).",
     "133 metas = 7 departamentos × 19 meses."),
    ("OTD-COM-05",
     "Plazo de entrega prometido por el proveedor: la columna existía pero estaba 100 % vacía.",
     "Ninguno — orden_compra.fecha_entrega_esperada ya existía.",
     "La creación de la orden de compra acepta y persiste la fecha prometida, validada contra la "
     "fecha de emisión.",
     "Campo «Fecha de entrega prometida» en el formulario de orden de compra.",
     "849 de 865 órdenes con fecha prometida; 825 pares promesa/llegada comparables."),
    ("OTD-COM-10",
     "Catálogo proveedor–producto: la tabla producto_proveedor existía con 0 filas, pese a que ya "
     "había líneas de compra reales.",
     "Ninguno — la tabla existía con costo, tiempo_entrega_dias, cantidad_minima y es_preferido.",
     "La recepción de compra alimenta automáticamente la oferta del proveedor (función SECURITY "
     "DEFINER), además del mantenimiento manual.",
     "Sección «Productos que ofrece» en la ficha del proveedor, con costo, plazo y marcador de "
     "preferido.",
     "1 106 ofertas de 11 proveedores; 1 093 con plazo; 1 040 preferidas."),
    ("OTD-INV-08",
     "Tope máximo de stock: la columna existía pero estaba 100 % vacía, así que el informe de "
     "sobre-stock devolvía cero filas.",
     "Ninguno — inventario.stock_maximo ya existía.",
     "El endpoint de niveles de inventario acepta y escribe stock_minimo y stock_maximo.",
     "Campos «Stock mínimo» y «Stock máximo» en el formulario de inventario.",
     "1 227 de 1 406 posiciones con tope máximo; 184 en sobre-stock hoy."),
    ("OTD-LOG-10",
     "Reembolsos sin registro transaccional: la tabla reembolso tenía 0 filas y el dato vivía "
     "disperso en la devolución.",
     "Ninguno — la tabla reembolso ya existía.",
     "La transición a 'reembolsada' de la devolución inserta además la fila en reembolso con "
     "monto, vía y fecha.",
     "Ninguno — la pantalla de reembolso del tablero de devoluciones ya capturaba monto y vía.",
     "85 reembolsos por $44 525,63."),
    ("OTD-LOG-11",
     "Costo real del envío: la columna era NOT NULL pero valía 0.00 en todos los envíos, porque "
     "nunca se calculaba desde la tarifa.",
     "Ninguno — envio.costo y peso_total_kg ya existían, y tarifa_envio tenía las tarifas.",
     "Al despachar, el costo se calcula desde la tarifa activa de la zona del pedido (la cadena "
     "dirección → zona → tarifa ya existía para asignar transportista) y se persiste el peso.",
     "Ninguno de captura (cálculo automático); el costo se muestra solo a roles con acceso a "
     "montos.",
     "2 848 de 2 872 envíos con costo, $32 723,25 en total."),
    ("OTD-SOP-02",
     "El tiempo prometido del ticket no era consultable: no había columna de fecha límite y el "
     "plazo por urgencia vivía solo en código del backend.",
     "Columna NUEVA ticket_soporte.fecha_limite, con permisos y RLS según el patrón del rol de "
     "soporte.",
     "El servicio de soporte la persiste al crear el ticket y la recalcula si se cambia la "
     "prioridad; el histórico se completó con el plazo derivado de su urgencia.",
     "Ninguno de captura — la bandeja ya mostraba «vence en / VENCIDO» y pasó a leer la columna.",
     "248 de 248 tickets con fecha límite; 76 cierres para contrastar."),
    ("OTD-SOP-08",
     "El reclamo no se podía ligar a un producto: el ticket solo referenciaba cliente y pedido.",
     "Columna NUEVA ticket_soporte.producto_variante_id, con permisos de columna sin dinero para "
     "el rol de soporte.",
     "El servicio de soporte escribe el producto al crear el ticket cuando el cliente lo indica.",
     "Selector opcional de producto en el formulario de creación del ticket.",
     "142 de 248 tickets ligados a un producto."),
    ("OTD-GER-09",
     "Intentos de entrada al sistema: la tabla log_acceso existía completa pero con 0 filas — el "
     "inicio de sesión nunca la escribía.",
     "Ninguno — la tabla existía (exitoso, motivo_fallo, ip_origen, email_intentado).",
     "El flujo de inicio de sesión inserta cada intento, exitoso o fallido, vía función SECURITY "
     "DEFINER en transacción propia, con correo intentado, IP y motivo.",
     "Ninguno — el formulario de inicio de sesión ya existía.",
     "1 537 intentos, 201 fallidos, 4 motivos, 19 meses."),
]
table(
    ["ID", "QUÉ FALTABA", "CAMBIO EN BASE DE DATOS", "CAMBIO EN BACKEND",
     "CAMBIO EN FORMULARIO / UI", "DATO HOY (verificado)"],
    brechas,
    widths=[0.7, 1.3, 1.3, 1.4, 1.2, 0.9],
    font_size=8,
)

p()
p("Las diez brechas comparten una lección que conviene explicitar: en siete de los diez casos "
  "NO faltaba nada en la base de datos — la columna o la tabla ya estaban declaradas y vacías. "
  "Lo que faltaba era el flujo que las escribiera y, en cuatro casos, el campo del formulario "
  "donde el usuario captura el dato. Un análisis táctico que se hubiera limitado a leer el "
  "esquema habría dado por cubiertos esos siete objetivos y habría producido siete informes "
  "vacíos.")

h(2, "5.1. Objetivos que aún esperan volumen de datos")
p("En estos dos casos el esquema está completo y el flujo ya escribe el dato: no hay cambio de "
  "base de datos, backend ni formulario. Lo que falta es densidad de operación antes de que el "
  "informe agregado discrimine algo.")
table(
    ["ID", "CONTEO REAL QUE LO LIMITA HOY", "POR QUÉ NO SE DA POR CUBIERTO"],
    [
        ("OTD-COM-09",
         "8 devoluciones a proveedor, 6 con resolución (3 reposiciones y 3 notas de crédito por "
         "$4 196,85), en 6 meses distintos.",
         "Un agregado «monto recuperado por proveedor y período» sobre 11 proveedores y 19 meses "
         "no discrimina con 6 observaciones."),
        ("OTD-GER-07",
         "123 líneas efectivamente promocionadas frente a 4 133 líneas de línea base de esos "
         "mismos productos.",
         "La ventana «durante» es demasiado pequeña frente a la base para sostener la comparación "
         "antes/durante producto por producto. Limitación aceptada y declarada."),
    ],
    widths=[0.9, 2.7, 2.9],
    font_size=9,
)

# ═══════════ 6. REPORTES SIMPLES IMPLEMENTADOS ════════════════════════════
doc.add_page_break()
h(1, "6. Reportes simples implementados")
p("La segunda parte del enunciado pide implementar los reportes simples. De los 29 objetivos "
  "clasificados como SIMPLES, 28 están construidos y consultables en la aplicación, cubriendo "
  "los SEIS departamentos. Se consultan POR PANTALLA, con filtros y registros visibles: no se "
  "entregan como PDF descargable — los PDF quedan reservados para los documentos operativos "
  "(facturas, guías de retorno, comprobantes).")
p("Los 28 informes comparten una sola arquitectura: la ruta "
  "GET /api/informes/{departamento}/{informe} devuelve siempre el mismo sobre de respuesta "
  "({items, total, page, size, resumen}), y una única pantalla genérica los pinta todos a partir "
  "de un archivo declarativo por departamento. Añadir un departamento completo cuesta dos clases "
  "de backend, un archivo de declaración y el enganche de navegación; ningún componente ni estilo "
  "nuevo. Toda consulta corre dentro de una transacción de solo lectura que asume el rol del "
  "usuario en el motor, de modo que los permisos por columna, la seguridad a nivel de fila y la "
  "restricción por horario los aplica PostgreSQL, no la aplicación.")

impl = [
    (VEN, "OTD-VEN-01", "ventas/cartera-pedidos", "estado, canal, desde/hasta, buscar",
     "4 083", "Vendedor, Gerente, Administrador"),
    (VEN, "OTD-VEN-02", "ventas/por-vendedor", "desde/hasta",
     "9 vendedores + canal en línea", "Gerente, Administrador (Vendedor: solo lo propio)"),
    (VEN, "OTD-VEN-08", "ventas/carritos-abandonados", "estado, días mínimos",
     "216 abandonados", "Gerente, Vendedor, Administrador"),
    (VEN, "OTD-VEN-10", "ventas/moderacion", "tipo, días mínimos",
     "53 reseñas + 13 preguntas", "Administrador, Gerente"),
    (VEN, "OTD-VEN-15", "ventas/avance-meta", "período (mes)",
     "133 metas", "Gerente, Vendedor, Administrador"),
    (COM, "OTD-COM-01", "compras/ordenes", "estado, proveedor, desde/hasta",
     "865", "Compras, Gerente, Administrador"),
    (COM, "OTD-COM-02", "compras/cuentas-por-pagar", "estado, situación, proveedor",
     "839 (276 vivas)", "Compras, Gerente, Administrador"),
    (COM, "OTD-COM-08", "compras/defectuosos", "estado, origen, proveedor, buscar",
     "38", "Compras, Gerente, Bodega (sin montos)"),
    (COM, "OTD-COM-10", "compras/catalogo-proveedor", "buscar, proveedor, oferta",
     "1 106", "Compras, Gerente"),
    (INV, "OTD-INV-01", "inventario/bajo-minimo", "bodega, buscar",
     "162 de 1 406", "Bodega, Compras, Gerente, Administrador"),
    (INV, "OTD-INV-02", "inventario/stock-bodega", "bodega, situación, buscar",
     "1 406", "Bodega, Compras, Vendedor, Gerente, Administrador"),
    (INV, "OTD-INV-03", "inventario/kardex", "buscar, bodega, naturaleza, tipo, desde/hasta",
     "13 287", "Bodega, Gerente, Administrador"),
    (INV, "OTD-INV-05", "inventario/ajustes", "tipo, estado, motivo, bodega, desde/hasta",
     "53", "Bodega, Gerente, Administrador"),
    (INV, "OTD-INV-06", "inventario/transferencias", "estado, bodega, desde/hasta",
     "71", "Bodega, Gerente, Administrador"),
    (INV, "OTD-INV-07", "inventario/valor-inventario", "bodega, categoría",
     "19 filas ($22,02 M)", "CON MONTO: Gerente, Administrador, Analista"),
    (INV, "OTD-INV-08", "inventario/sobre-stock", "bodega, buscar",
     "184", "Bodega, Compras, Gerente"),
    (LOG, "OTD-LOG-01", "logistica/cola-despacho", "estado, canal, transportista, buscar",
     "48", "Despacho, Gerente, Administrador"),
    (LOG, "OTD-LOG-02", "logistica/envios", "estado, transportista, desde/hasta, buscar",
     "2 872", "Despacho, Gerente, Administrador"),
    (LOG, "OTD-LOG-06", "logistica/devoluciones", "estado, motivo, desde/hasta, buscar",
     "196", "Despacho, Soporte, Gerente, Bodega (sin montos)"),
    (LOG, "OTD-LOG-11", "logistica/costo-envio", "zona, transportista, desde/hasta",
     "9 filas ($32 723,25)", "CON MONTO: Gerente, Administrador"),
    (SOP, "OTD-SOP-01", "soporte/bandeja", "estado, prioridad, categoría, agente, buscar",
     "248 (128 vivos)", "Soporte, Gerente, Administrador"),
    (SOP, "OTD-SOP-04", "soporte/por-categoria", "desde/hasta",
     "8 categorías", "Soporte, Gerente"),
    (SOP, "OTD-SOP-05", "soporte/por-agente", "desde/hasta",
     "7 filas", "Soporte, Gerente, Administrador"),
    (GER, "OTD-GER-01", "gerencia/foto-dia", "fecha",
     "20 filas agregadas", "Gerente, Administrador"),
    (GER, "OTD-GER-04", "gerencia/cupones", "situación, tipo, buscar",
     "33 (7 vigentes)", "Gerente, Administrador"),
    (GER, "OTD-GER-06", "gerencia/marketing", "tipo, vigencia, buscar",
     "65 (20 vigentes)", "Gerente, Administrador"),
    (GER, "OTD-GER-08", "gerencia/auditoria", "usuario, tabla, acción, desde/hasta",
     "7 073", "SENSIBLE: Administrador, Gerente"),
    (GER, "OTD-GER-09", "gerencia/accesos", "resultado, correo, desde/hasta",
     "1 537", "SENSIBLE: Administrador, Gerente"),
]
assert len(impl) == 28
_ids_impl = {o[1] for o in OBJ if o[8]}
assert {f[1] for f in impl} == _ids_impl, "La tabla de reportes no coincide con el catálogo"

t = table(
    ["DEPTO.", "ID", "RUTA DEL INFORME", "FILTRO PRINCIPAL", "FILAS HOY",
     "ROL DESTINATARIO"],
    impl,
    widths=[0.8, 0.9, 1.5, 1.6, 0.8, 1.5],
    font_size=8,
)
merge_first_col(t, impl, font_size=8)

p()
h(2, "6.1. Segregación financiera de los reportes")
p("La regla del sistema es que Bodega y Despacho nunca ven montos de dinero: sus tableros "
  "trabajan con cantidades, estados y fechas, y las vistas valorizadas de esos mismos fenómenos "
  "se muestran a Gerencia, Administración y Analista. Implementar los 28 informes obligó a "
  "descubrir que esa regla no se sostiene en un único lugar, sino en TRES, según qué capa pueda "
  "ser la última línea de defensa:")
for txt in [
    "El MOTOR la respalda cuando el rol simplemente no tiene permiso de lectura sobre la columna "
    "de dinero. Es el caso de los informes de Ventas: el grupo de bodega no puede leer el total "
    "del pedido, así que aunque alguien escribiera la consulta, PostgreSQL la rechazaría.",
    "La RUTA la impone cuando el motor no puede. Bodega SÍ puede leer el costo de la variante "
    "—lo necesita para valorizar su kardex al recibir mercancía—, de modo que el informe de valor "
    "del inventario (OTD-INV-07) no se protege escondiendo columnas sino cerrando su propio "
    "endpoint. Lo mismo ocurre con el costo de envío (OTD-LOG-11) y con la auditoría (OTD-GER-08), "
    "que el rol de analista sí puede leer en la base.",
    "La CONSULTA la aplica cuando un mismo informe se comparte entre roles con y sin acceso a "
    "dinero. En el informe de artículos defectuosos (OTD-COM-08) entra Bodega, y el motor no lo "
    "impide porque el flujo operativo de la devolución al proveedor necesita esos costos: la "
    "barrera es que el SQL no selecciona ninguna columna de monto.",
]:
    doc.add_paragraph(txt, style="List Bullet")
p("La matriz de autorización se verificó llamando cada endpoint con cada uno de los ocho roles "
  "del sistema. Los resultados confirman el diseño: Despacho recibe respuesta en los informes de "
  "logística sin dinero y 403 en el de costo de envío; Bodega entra a los seis informes de "
  "cantidades de inventario y recibe 403 en el de valor del inventario; los informes de "
  "auditoría y de intentos de acceso responden únicamente a Administrador y Gerente. Un valor de "
  "filtro fuera de la lista blanca devuelve 400 con el listado de permitidos, y cualquier método "
  "distinto de GET sobre una ruta de informes devuelve 403: un informe nunca escribe.")

h(2, "6.2. Visualización y destinatarios de los 68 objetivos")
p("Para cada objetivo del catálogo —implementado o no— quedó definida su visualización y los "
  "roles que la consultan. Los marcados «(ETL)» corresponden a informes compuestos, cuya "
  "visualización se construirá sobre ClickHouse.")

filas_dash = [
    (o[0], o[1], o[6], o[7], "Implementado" if o[8] else ("Pendiente" if o[3] else "ETL"))
    for o in OBJ
]
t = table(
    ["DEPTO.", "ID", "VISUALIZACIÓN (pantalla con filtros)", "ROLES QUE LO VEN", "ESTADO"],
    filas_dash,
    widths=[0.8, 1.0, 2.3, 1.9, 0.8],
    font_size=8,
)
merge_first_col(t, filas_dash, font_size=8)

# ═══════════ 7. ANÁLISIS DE RESULTADOS ════════════════════════════════════
doc.add_page_break()
h(1, "7. Análisis de resultados")
table(
    ["CLASIFICACIÓN", "FUENTE", "CANTIDAD", "PORCENTAJE"],
    [
        ("Informes simples", BDR, str(N_SIMPLES), "42,6 %"),
        ("Informes compuestos", BDC, str(N_COMPUESTOS), "57,4 %"),
        ("Total", "", "68", "100 %"),
    ],
    widths=[1.9, 2.3, 1.2, 1.4],
    font_size=10,
)
p()
filas_dep = [
    ("Ventas", VEN), ("Compras", COM), ("Inventario / Bodega", INV),
    ("Logística / Despacho", LOG), ("Soporte", SOP),
    ("Gerencia / Dirección (incl. Marketing)", GER),
]
rows_dep = []
for etiqueta, d in filas_dep:
    tot, s, c, i = RESUMEN_DEPTO[d]
    rows_dep.append((etiqueta, str(tot), str(s), str(c), "%d de %d" % (i, s)))
rows_dep.append(("Total", "68", str(N_SIMPLES), str(N_COMPUESTOS),
                 "%d de %d" % (N_IMPL, N_SIMPLES)))
table(
    ["DEPARTAMENTO", "OBJETIVOS", "SIMPLES", "COMPUESTOS", "REPORTES IMPLEMENTADOS"],
    rows_dep,
    widths=[2.3, 1.0, 1.0, 1.2, 1.5],
    font_size=10,
)
p()
p("Verificación aritmética: 15+12+10+12+8+11 = 68 objetivos; 5+5+7+4+3+5 = 29 simples; "
  "10+7+3+8+5+6 = 39 compuestos; 29+39 = 68. Reportes implementados 5+4+7+4+3+5 = 28 de 29 "
  "simples. Por estado de factibilidad: 66 factibles hoy + 0 que requieren cambio en el sistema "
  "+ 2 que requieren volumen de datos = 68.")
p("De los 68 informes tácticos, 29 son simples y se resuelven directamente sobre la BDR "
  "PostgreSQL —28 ya construidos y consultables—, mientras que 39 son compuestos y se procesarán "
  "en la base columnar ClickHouse. Los compuestos comparten tres rasgos que desaconsejan "
  "ejecutarlos sobre la base transaccional: (a) barren volúmenes históricos completos en lugar de "
  "filas puntuales — el kardex ya tiene 13 287 movimientos y el historial de estados de pedido "
  "24 608 transiciones; (b) exigen agregaciones (sumas, promedios, tasas) y agrupaciones por "
  "dimensiones de tiempo, categoría, proveedor o zona; y (c) su patrón de lectura masiva por "
  "columnas es exactamente el que un motor columnar optimiza. Ejecutarlos sobre PostgreSQL "
  "competiría por recursos con las transacciones operativas (ventas, checkout, kardex) y "
  "degradaría el sistema en producción; por eso los 39 compuestos justifican el pipeline ETL "
  "hacia ClickHouse y su orquestación con Apache Airflow.")
p("La asimetría entre departamentos es deliberada y responde a la necesidad real de cada área, "
  "no a una cuota. Ventas concentra el mayor número de objetivos porque concentra el ingreso; "
  "Compras le sigue de cerca porque en una distribuidora mayorista el abastecimiento es el centro "
  "de costo dominante — compra cuatro veces lo que vende en el período analizado — y su control "
  "de deuda, plazos y cumplimiento del proveedor es donde se gana el margen; Logística carga "
  "además todo el ciclo de devoluciones de clientes; Inventario combina el control del presente "
  "(existencias, kardex, ajustes — de ahí su mayoría de simples, siete de diez) con la evolución "
  "del capital almacenado; y Soporte, un área más pequeña, tiene un catálogo proporcional a lo "
  "que dirige.")
p("También es revelador que 9 de los 11 objetivos incorporados en la ampliación del catálogo "
  "resultaran compuestos: lo que faltaba no era tanto ver el presente — eso el catálogo ya lo "
  "cubría casi por completo — sino medir tiempos de ciclo, comparar períodos y consolidar dinero, "
  "exactamente el tipo de pregunta que justifica la base columnar. Y en sentido inverso, el único "
  "objetivo que cambió de clasificación lo hizo al implementarse: el costo de envío por zona y "
  "transportista se reveló como una foto agregada del presente, no como una serie histórica, "
  "confirmando que la regla de clasificación se sostiene mejor cuando se contrasta contra el "
  "informe realmente construido y no contra la intuición inicial.")

# ═══════════ 8. CONCLUSIÓN ════════════════════════════════════════════════
doc.add_page_break()
h(1, "8. Conclusión")
p("El análisis del nivel táctico de RetailMind demuestra que las necesidades de dirección y "
  "control de los seis departamentos se cubren con 68 objetivos tácticos verificados uno a uno "
  "contra la base de datos real del sistema (110 tablas): 29 informes simples, atendidos por la "
  "base relacional PostgreSQL que ya soporta la operación, y 39 compuestos, que por su naturaleza "
  "histórica y agregada se destinan a la base columnar ClickHouse.")
p("La verificación honesta contra los datos reales — no contra el esquema en papel — permitió "
  "clasificar cada objetivo por factibilidad y, sobre todo, detectar diez brechas donde la "
  "columna existía pero nadie la escribía. Las diez se cerraron aplicando el método completo en "
  "sus tres capas: base de datos, backend y formulario. Hoy no queda ninguna brecha de sistema "
  "abierta: 66 de los 68 objetivos son respondibles con el dato real, y los 2 restantes solo "
  "esperan que la operación densifique su histórico, sin exigir cambio alguno al sistema.")
p("La segunda parte del encargo también está cumplida: 28 de los 29 informes simples están "
  "implementados y se consultan por pantalla en los seis departamentos, con filtros, paginación "
  "del lado del servidor e indicadores de cabecera, bajo un contrato único que permitió "
  "construirlos con una sola pantalla genérica. La segregación financiera del sistema se "
  "respeta en las tres capas que pueden imponerla — el motor, la ruta y la propia consulta —, "
  "verificada llamando cada endpoint con cada uno de los ocho roles.")
p("La arquitectura híbrida existente resulta, por tanto, adecuada para el nivel táctico: la BDR "
  "responde el «ahora» de cada área y la BD columnar responderá el «cómo evoluciona». El "
  "siguiente paso natural, y el único que queda del nivel táctico, es la construcción del "
  "pipeline ETL hacia ClickHouse orquestado con Apache Airflow — programación, dependencias "
  "entre tareas, reintentos y monitoreo — para automatizar la alimentación de los 39 informes "
  "compuestos.")

import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "RetailMind_T11_Analisis_Tactico.docx")
doc.save(out)
print("Guardado:", out)
print("Objetivos:", len(OBJ), "| Simples:", N_SIMPLES, "| Compuestos:", N_COMPUESTOS)
print("Factibilidad -> FACTIBLE HOY:", N_FH, "| REQUIERE CAMBIO:", N_RC,
      "| REQUIERE VOLUMEN:", N_RV)
print("Reportes simples implementados:", N_IMPL, "de", N_SIMPLES)
