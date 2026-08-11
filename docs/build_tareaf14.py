# -*- coding: utf-8 -*-
"""
Generador de docs/RetailMind_TareaF14_Arquitectura_Por_Nivel.docx

TareaF 14 (formativa): 1) OPERATIVO COMPLETO  2) TACTICO COMPLETO
                       3) ESTRATEGICO AGG

REGLA DE ESTE DOCUMENTO: toda cifra que aparece abajo fue consultada contra el
sistema REAL el 2026-08-09 (contenedores `retailmind-postgres-1`,
`retailmind-clickhouse-1` y la API en el 8080). No se copio ni un numero de
CLAUDE.md ni de ningun documento previo sin volver a medirlo; donde la
documentacion existente discrepaba del sistema, manda el sistema y la
discrepancia se declara en la seccion 7 del documento.

Salida: SOLO el .docx (el PDF, si hace falta, lo produce el usuario).
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Estilos base (mismos que build_t11.py, para uniformidad de entregables) ──
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def _set_heading_color(style_name, size, rgb):
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor(*rgb)
    st.font.bold = True


_set_heading_color("Heading 1", 18, (0x1F, 0x3B, 0x5B))
_set_heading_color("Heading 2", 14, (0x1F, 0x3B, 0x5B))
_set_heading_color("Heading 3", 12, (0x2E, 0x5A, 0x88))

CENTER = WD_ALIGN_PARAGRAPH.CENTER
JUST = WD_ALIGN_PARAGRAPH.JUSTIFY


def h(level, text):
    return doc.add_heading(text, level=level)


def p(text="", bold=False, italic=False, size=None, align=JUST):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    par.alignment = align
    return par


def bullet(text, bold_prefix=None):
    par = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = par.add_run(bold_prefix)
        r.bold = True
    par.add_run(text)
    par.alignment = JUST
    return par


def table(headers, rows, widths=None, font_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run("" if val is None else str(val))
            r.font.size = Pt(font_size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ══════════════════════════ PORTADA ══════════════════════════
h(0, "RetailMind — Arquitectura por nivel")
p("Universidad Técnica Estatal de Quevedo (UTEQ) · Facultad de Ciencias de la Ingeniería",
  bold=True, align=CENTER)
p("Construcción de Software — sexto semestre", align=CENTER)
p("TareaF 14 (formativa): 1) Operativo completo · 2) Táctico completo · 3) Estratégico AGG",
  align=CENTER)
p("Fecha del documento: 9 de agosto de 2026", align=CENTER)
p()
p("Nota de verificación. Todas las cifras de este documento se consultaron contra el "
  "sistema en funcionamiento el 9 de agosto de 2026: la base transaccional PostgreSQL 18.4 "
  "del contenedor (puerto 5432), el almacén analítico ClickHouse y la propia API en el "
  "puerto 8080. Ninguna cifra se copió de documentación anterior sin volver a medirla. Las "
  "diferencias encontradas entre la documentación del proyecto y el sistema real se declaran "
  "en la sección 7, no se ocultan.", italic=True, size=9)

doc.add_page_break()

# ══════════════════════════ 1. INTRODUCCIÓN ══════════════════════════
h(1, "1. Introducción: qué es RetailMind")

p("RetailMind es el proyecto capstone de la asignatura. Simula una empresa de comercio "
  "minorista multicanal de ticket alto radicada en Quevedo (Los Ríos, Ecuador): una tienda "
  "que vende por tres vías —mostrador, teléfono y tienda en línea— y que, por detrás, compra "
  "a proveedores, almacena, despacha, cobra, factura y atiende reclamos.")

p("El sistema no es una maqueta de pantallas. Es una operación completa sobre una base "
  "transaccional real de 111 tablas, con 19 meses de historia sembrada, sobre la que se "
  "construyeron después dos niveles de lectura: un nivel táctico de informes para los jefes "
  "de área y un nivel estratégico de tableros y modelos para la dirección. Los tres niveles "
  "que pide esta tarea son exactamente esas tres capas, y se sostienen sobre dos motores "
  "distintos y deliberadamente separados:")

bullet("es la única base transaccional. Todo lo que el negocio escribe —pedidos, "
       "inventario, facturas, tickets— vive aquí, y aquí vive también la seguridad "
       "(roles, políticas por fila, permisos por columna).",
       bold_prefix="PostgreSQL 18.4 ")
bullet("es solo analítica. Contiene el almacén de datos y la serie de eventos de "
       "navegación. Es una capa de LECTURA: si se apaga, la operación sigue funcionando "
       "entera y solo se degradan los informes compuestos y los tableros, con aviso en "
       "pantalla.",
       bold_prefix="ClickHouse ")

p("Esa separación es el invariante de diseño del proyecto y explica por qué los tres niveles "
  "se comportan distinto ante una caída: el operativo no depende de la analítica, y la "
  "analítica nunca escribe en el operativo.")

h(2, "1.1 Estado del despliegue el día de la medición")
p("El sistema está contenerizado. El archivo de composición declara nueve servicios "
  "repartidos en tres perfiles; el perfil por defecto levanta cuatro, que son los que "
  "estaban en marcha y sanos durante toda la verificación:")
table(
    ["Servicio", "Papel", "Estado medido"],
    [
        ["postgres", "Base transaccional (PostgreSQL 18.4), puerto 5432", "Activo y sano"],
        ["clickhouse", "Almacén analítico y eventos de navegación", "Activo y sano"],
        ["backend", "API Spring Boot 3.5 (Java 17), puerto 8080", "Activo y sano"],
        ["frontend", "Angular 17, puerto 4200", "Activo y sano"],
        ["etl · pgadmin", "Perfil «tools», se invocan a demanda", "No levantados"],
        ["airflow (3 servicios)", "Perfil «airflow», orquestación del ETL",
         "No levantados; el DAG corrió la noche anterior"],
    ],
    widths=[1.7, 3.2, 1.8])

p("El chequeo de salud de la API respondió «UP» con sus cuatro componentes —PostgreSQL, "
  "ClickHouse, base y entorno de Python— también en «UP».", size=10, italic=True)

# ══════════════════════════ 2. OPERATIVO ══════════════════════════
doc.add_page_break()
h(1, "2. Nivel OPERATIVO")

p("El nivel operativo es el sistema que la empresa usa para trabajar: registra hechos y los "
  "hace avanzar por un proceso. Su criterio de completitud no es «cuántas pantallas hay», "
  "sino si cada ciclo de negocio llega hasta el final y si el paso indebido está impedido.")

h(2, "2.1 Ciclos de negocio implementados y hasta dónde llegan")

table(
    ["Ciclo", "Recorrido implementado", "Evidencia medida (9-ago-2026)"],
    [
        ["Compra a proveedor",
         "Orden → aprobación de Gerencia/Administración → recepción (con rechazo en "
         "puerta) → factura → cuenta por pagar → pago",
         "868 órdenes; 839 facturas de compra; $22.467.387,27 facturados"],
        ["Inventario",
         "Entradas y salidas con kárdex encadenado, transferencias entre bodegas y ajustes "
         "con anulación por contramovimiento",
         "13.289 movimientos de kárdex; 1.406 posiciones; 133.226 unidades en stock; "
         "71 transferencias; 53 ajustes"],
        ["Venta",
         "Pedido → cobro → facturación → preparación (bodega) → despacho → entrega, con "
         "devolución solo después de entregar",
         "4.083 pedidos (2.213 web / 1.030 mostrador / 840 teléfono); 3.887 facturas; "
         "4.079 pagos; 24.610 hitos de historial"],
        ["Tienda del cliente",
         "Catálogo, carrito, lista de deseos, perfil y direcciones, pago simulado y «Mis "
         "pedidos»; el pedido en línea entra al MISMO ciclo de venta",
         "1.214 productos / 1.221 variantes; 290 carritos; 72 clientes"],
        ["Última milla",
         "Envío con transportista asignado por zona, seguimiento e incidencias de entrega "
         "con reprogramación o retorno al almacén",
         "2.872 envíos; 177 novedades de envío"],
        ["Logística inversa (RMA)",
         "Solicitud del cliente → revisión → guía de retorno → tránsito → recepción → "
         "inspección por ítem → reembolso → cierre",
         "197 devoluciones de cliente; 108 pedidos en estado «devuelto»"],
        ["Devolución a proveedor",
         "Pool de ítems defectuosos (de RMA o de recepción) → agrupación por proveedor → "
         "nota de crédito o reposición",
         "38 ítems defectuosos; 8 devoluciones a proveedor"],
        ["Soporte",
         "Ticket con número y prioridad automática por categoría, SLA, asignación, "
         "transiciones con guardias y reapertura",
         "249 tickets"],
        ["Marketing y descuentos",
         "Cupones y promociones que descuentan de verdad sobre el pedido y se prorratean "
         "en la factura",
         "33 cupones con 564 canjes; 24 promociones; 133 metas de venta"],
        ["Reseñas y preguntas",
         "Reseña con compra verificada y moderación",
         "344 reseñas"],
    ],
    widths=[1.35, 2.7, 2.65])

p("El estado de los 4.083 pedidos confirma que el ciclo se recorre hasta el final y no se "
  "queda a mitad: 3.592 entregados, 159 cancelados, 121 no entregados y 108 devueltos; solo "
  "103 pedidos están repartidos entre los seis estados intermedios (confirmado, pagado, "
  "facturado, en preparación, preparado y despachado). El catálogo de estados tiene once "
  "entradas.")

h(2, "2.2 Las compuertas: dónde se impone cada regla")

p("Una compuerta es la prohibición de saltarse un paso —facturar sin haber cobrado, "
  "despachar sin haber preparado—. En este sistema no viven todas en el mismo sitio, y esa "
  "es una decisión de diseño que conviene explicar, porque es lo que distingue una regla que "
  "se puede burlar de una que no.")

table(
    ["Se impone en…", "Qué garantiza", "Cuánto hay (medido)"],
    [
        ["EL MOTOR — restricciones CHECK",
         "Que un estado inválido no llegue a existir. Devolución, envío y orden de compra "
         "tienen su lista de estados fijada en la tabla; el canal del pedido solo admite "
         "web, tienda o teléfono",
         "Restricciones de estado en devolución (9 valores), envío (6) y orden de compra "
         "(6)"],
        ["EL MOTOR — disparadores",
         "Que los totales no se puedan escribir a mano (los recalcula el motor desde el "
         "detalle), que la ecuación del kárdex cuadre y que el uso de un cupón se cuente "
         "aunque haya concurrencia",
         "92 disparadores propios: 34 de horario, 6 de recálculo de totales, 2 de "
         "validación del kárdex, 2 de uso de cupón y el resto de sello de fecha"],
        ["EL MOTOR — permisos y políticas",
         "Que un rol no lea ni escriba lo que no le toca, aunque la aplicación se equivoque "
         "o alguien se conecte por fuera de ella",
         "95 políticas de seguridad por fila sobre 50 tablas; 109 columnas con permiso "
         "propio en 14 tablas"],
        ["LA APLICACIÓN — guardias de proceso",
         "El ORDEN de los pasos y la idempotencia: sin aprobar no se recibe, sin recibir "
         "completo no se factura, sin factura no hay pago. Responden con un mensaje que "
         "explica el porqué",
         "84 guardias de estado: ventas 30, compras 18, devoluciones 14, soporte 12, "
         "carrito 8, inventario 2"],
    ],
    widths=[1.7, 3.1, 1.9])

p("El reparto es honesto y conviene decirlo tal cual: el orden del proceso lo impone la "
  "APLICACIÓN, no el motor. Lo que el motor garantiza es que los datos no queden "
  "incoherentes —totales, kárdex, estados existentes— y que nadie vea ni toque lo que no le "
  "corresponde. Una transición fuera de orden ejecutada directamente contra la base, "
  "saltándose la API, no sería rechazada por un disparador de secuencia, porque ese "
  "disparador no existe: la máquina de estados vive en Java.")

h(2, "2.3 Seguridad a nivel de base de datos")

p("Es la parte más distintiva del proyecto. La aplicación se conecta con un usuario técnico "
  "sin privilegios de negocio y, dentro de cada transacción, asume el rol del usuario "
  "autenticado; el nombre del rol viaja como parámetro ligado y nunca concatenado en el SQL. "
  "La consecuencia práctica es que todo acceso a la base debe ir dentro de una transacción: "
  "hay 285 anotaciones transaccionales en el código por esa razón.")

table(
    ["Mecanismo", "Cifra verificada", "Qué significa"],
    [
        ["Roles de grupo en el motor", "9",
         "administrador, gerente, vendedor, compras, bodega, despacho, cliente, analista y "
         "soporte"],
        ["Permisos concedidos a esos 9 roles", "1.355",
         "De 2.192 concesiones de tabla en total; el resto pertenece al propietario y a los "
         "dos usuarios técnicos (aplicación y ETL)"],
        ["Políticas de seguridad por fila", "95, sobre 50 tablas",
         "50 son de restricción horaria; las otras 45 aíslan a cada cliente a lo suyo"],
        ["Tablas con seguridad por fila activada", "50",
         "Coincide con las que tienen política: no hay tabla protegida sin regla que la "
         "acompañe"],
        ["Disparadores de restricción horaria", "34",
         "Bloquean la escritura fuera de la ventana horaria del rol"],
        ["Ventanas horarias configuradas", "56 filas, 8 roles",
         "Las 56 están hoy en 00:00–24:00, es decir 24/7: el mecanismo está intacto, lo que "
         "se abrió son los datos"],
        ["Columnas con permiso propio", "109, en 14 tablas",
         "Es la segregación financiera: bodega y despacho no leen importes"],
        ["Funciones con privilegio elevado", "17",
         "Recálculos de totales y operaciones administrativas que deben correr por encima "
         "del rol que las invoca"],
        ["Roles personalizados creados", "0",
         "La capacidad existe —hay pantalla de administración—, pero hoy no hay ninguno en "
         "uso"],
    ],
    widths=[2.0, 1.35, 3.35])

p("Sobre la restricción horaria conviene ser preciso, porque se presta a malentendido: hoy "
  "no bloquea a nadie —las 56 ventanas cubren la semana completa—, pero no se eliminó ni se "
  "debilitó. Los 34 disparadores, las 50 políticas y las funciones que evalúan el horario "
  "siguen en su sitio; lo que cambió fueron las filas de configuración. Basta volver a "
  "estrechar una ventana para que el bloqueo reaparezca.")

h(2, "2.4 Tamaño del nivel operativo")
table(
    ["Magnitud", "Medida"],
    [
        ["Tablas en el esquema público de PostgreSQL", "111"],
        ["Índices · claves foráneas · funciones", "381 · 198 · 61"],
        ["Tamaño de la base transaccional", "65 MB"],
        ["Clases Java · controladores · endpoints REST", "155 · 46 · 361"],
        ["Rutas declaradas en el frontend Angular", "67"],
        ["Usuarios registrados (72 clientes y 17 de personal)", "89"],
        ["Registros de auditoría · de acceso", "7.540 · 2.433"],
    ],
    widths=[4.0, 2.7])

# ══════════════════════════ 3. TÁCTICO ══════════════════════════
doc.add_page_break()
h(1, "3. Nivel TÁCTICO")

p("El nivel táctico responde preguntas de jefe de área: cómo va mi cartera, quién vende "
  "cuánto, qué le debo a cada proveedor, cuánto tarda un envío. Son informes de consulta por "
  "pantalla, sin PDF, organizados por departamento.")

h(2, "3.1 El catálogo de objetivos")

p("El catálogo declara 69 objetivos tácticos repartidos en seis departamentos, más uno "
  "descartado por medición: la segmentación mayorista/minorista, que se comprobó que no "
  "existe en los datos. De esos 69, treinta son SIMPLES y treinta y nueve COMPUESTOS. La "
  "diferencia no es de dificultad sino de FUENTE, y determina de qué motor lee cada uno:")

table(
    ["Tipo", "Qué pregunta responde", "De dónde lee", "Cuántos"],
    [
        ["SIMPLE", "La foto de HOY: un estado presente, un listado, un agregado del período",
         "PostgreSQL, la base transaccional, en vivo", "30"],
        ["COMPUESTO",
         "El RECORRIDO: evolución mes a mes, comparación entre períodos, series históricas",
         "ClickHouse, el almacén «retailmind_dwh»", "39"],
    ],
    widths=[1.1, 2.7, 2.2, 0.7])

p("El reparto por departamento de esos 69 objetivos es: Ventas 17, Compras 12, Logística 12, "
  "Gerencia 11, Inventario 10 y Soporte 8.")

h(2, "3.2 Qué hay construido, medido ruta por ruta")

p("No basta con contar objetivos en un documento: se llamó a cada ruta de la API con un "
  "usuario administrador y se registró la respuesta. El resultado:")

table(
    ["Departamento", "Rutas simples", "Rutas compuestas", "Total"],
    [
        ["Ventas", "6", "11", "17"],
        ["Inventario", "7", "3", "10"],
        ["Compras", "5", "8", "13"],
        ["Logística", "4", "9", "13"],
        ["Soporte", "3", "5", "8"],
        ["Gerencia", "5", "7", "12"],
        ["TOTAL", "30", "43", "73"],
    ],
    widths=[2.2, 1.5, 1.6, 1.4])

p("Las 43 rutas compuestas no son 43 objetivos: son los 39 objetivos compuestos más cuatro "
  "extras —los dos modelos estratégicos, la serie mensual del costo de envío (declarada "
  "fuera del catálogo) y la previsión de demanda, que se sirve en dos departamentos—. Es una "
  "distinción que conviene mantener: contar rutas y llamarlas objetivos infla la cifra en "
  "un 10 %.")

p("Resultado de la prueba: las 43 rutas compuestas respondieron correctamente, todas juntas "
  "en 3,35 segundos, 78 milisegundos de media. De las 30 simples, 29 respondieron con sus "
  "parámetros por defecto y una devolvió un aviso controlado: el informe de avance sobre la "
  "meta contesta «no hay meta de ventas vigente para 8/2026», porque las metas sembradas "
  "llegan hasta julio de 2026. Consultado con un mes que sí tiene meta, funciona y devuelve "
  "$235.000,00 de meta contra $225.463,58 vendidos, un 95,94 % de avance. Es una limitación "
  "del dato, no un fallo del informe, y el mensaje al usuario lo dice con esas palabras.")

h(2, "3.3 Cómo se validan los números de este nivel")

p("Los informes simples leen la base transaccional, así que su número es el de la operación "
  "por definición. Los compuestos leen el almacén, que es una COPIA, y ahí el riesgo real es "
  "que la copia se desvíe sin que nadie lo note. Contra eso hay dos redes:")

bullet("cada tabla del almacén se carga primero en una tabla temporal, se compara contra "
       "PostgreSQL y solo si cuadra se publica intercambiándola con la anterior. Si la "
       "validación falla, la tabla publicada no se toca y el informe sigue mostrando la "
       "carga anterior en lugar de una cifra falsa.",
       bold_prefix="Carga atómica. ")
bullet("un verificador independiente con 49 controles que cruzan el almacén contra "
       "PostgreSQL al centavo, tres de ellos cruzando tablas entre sí dentro del propio "
       "almacén.", bold_prefix="Cuarenta y nueve controles. ")

p("Para este documento no se ejecutó ese verificador —habría escrito una fila en la bitácora "
  "del ETL y la sesión era de solo lectura—. En su lugar se repitieron a mano cinco de sus "
  "comprobaciones, consultando los dos motores por separado:")

table(
    ["Control", "PostgreSQL", "ClickHouse (almacén)", "Diferencia"],
    [
        ["Pedidos no cancelados", "3.924", "3.924", "0"],
        ["Venta de esos pedidos", "$5.498.570,35", "$5.498.570,35", "$0,00"],
        ["Unidades vendidas en línea de pedido", "20.687", "20.687", "0"],
        ["Movimientos de kárdex", "13.289", "13.289", "0"],
        ["Cierre de stock del último mes frente al inventario vivo", "133.226 uds",
         "133.226 uds", "0"],
    ],
    widths=[2.4, 1.5, 1.7, 1.1])

p("El último control es el más exigente de los cinco: el almacén reconstruye el saldo de "
  "inventario mes a mes desde el kárdex, y su último mes coincide exactamente con el stock "
  "que hoy tiene la operación.", italic=True, size=10)

# ══════════════════════════ 4. ESTRATÉGICO ══════════════════════════
doc.add_page_break()
h(1, "4. Nivel ESTRATÉGICO")

h(2, "4.1 Misión, visión y objetivos estratégicos")

p("Misión. «Proveer a nuestros clientes acceso a un catálogo diverso de productos a precios "
  "altamente competitivos, eliminando intermediarios innecesarios. Operamos mediante "
  "procesos ágiles y confiables de abastecimiento, control de inventario y logística de "
  "entrega, garantizando la máxima disponibilidad y el cumplimiento exacto en cada pedido.»",
  italic=True)

p("Visión. «Para el año 2030, ser la red de distribución comercial referente en la Costa "
  "ecuatoriana, reconocida por conectar eficientemente a proveedores y consumidores. "
  "Lideraremos el mercado mediante una operación confiable y la toma de decisiones "
  "estratégicas basadas en datos, anticipando la demanda y optimizando cada eslabón de "
  "nuestra cadena de abastecimiento.»", italic=True)

p("La base estratégica vigente fija seis objetivos. Se numeran desde OE-06 porque los "
  "identificadores OE-01 a OE-05 estaban tomados por dos juegos anteriores incompatibles "
  "entre sí, y reutilizarlos habría roto las citas de los documentos ya entregados.")

table(
    ["ID", "Nombre", "Tablero que lo sirve"],
    [
        ["OE-06", "Consolidación de la experiencia omnicanal", "T-1"],
        ["OE-07", "Rentabilidad por volumen y rotación", "T-2"],
        ["OE-08", "Fidelización y retención de clientes", "T-3"],
        ["OE-09", "Eficiencia operativa", "T-4 y T-5"],
        ["OE-10", "Liderazgo en decisiones basadas en datos", "T-7"],
        ["OE-11", "Excelencia en la cadena de abastecimiento", "T-6"],
    ],
    widths=[0.7, 3.4, 2.6])

h(2, "4.2 Los siete tableros de dirección")

p("Cada tablero responde en UNA sola llamada: devuelve sus indicadores, sus bloques, sus "
  "salvedades y la fecha del dato. Se hizo así a propósito, en lugar de una llamada por "
  "elemento, para que todos compartan filtros y se degraden a la vez; con seis peticiones, "
  "una caída a mitad de carga dejaría medio tablero pintado y creíble.")

p("Los siete cubren las 19 decisiones de dirección declaradas en el diseño —verificado en el "
  "código: los siete servicios declaran 19 identificadores de decisión distintos—. Todos "
  "respondieron correctamente y se midieron tres veces cada uno:")

table(
    ["Tablero", "Objetivo", "Tablas del almacén que lee", "Tiempo (3 medidas)"],
    [
        ["T-1 Omnicanal", "OE-06", "fact_pedido, fact_venta_linea, fact_flujo_caja",
         "0,20 / 0,17 / 0,12 s"],
        ["T-2 Rentabilidad y rotación", "OE-07",
         "fact_venta_linea, fact_stock_mensual, fact_movimiento_inventario, dim_producto",
         "0,23 / 0,18 / 0,14 s"],
        ["T-3 Cliente y posventa", "OE-08",
         "fact_pedido, fact_devolucion, fact_devolucion_linea, fact_ticket, fact_resena, "
         "dim_cliente, dim_producto", "0,19 / 0,19 / 0,17 s"],
        ["T-4 Operación y última milla", "OE-09",
         "fact_envio, fact_novedad_envio, fact_pedido, fact_devolucion, "
         "fact_movimiento_inventario", "0,20 / 0,15 / 0,13 s"],
        ["T-5 Costo de la operación", "OE-09", "fact_envio, fact_pedido, fact_devolucion",
         "0,13 / 0,12 / 0,12 s"],
        ["T-6 Abastecimiento", "OE-11",
         "fact_orden_compra, fact_compra_linea, fact_flujo_caja, fact_devolucion_proveedor",
         "0,20 / 0,18 / 0,17 s"],
        ["T-7 Gobierno del dato", "OE-10",
         "las 19 tablas del modelo más la bitácora del ETL", "0,22 / 0,20 / 0,19 s"],
    ],
    widths=[1.5, 0.7, 2.9, 1.2])

p("El tablero más lento del sistema tarda 227 milisegundos en su primera consulta en frío. "
  "Esa cifra es la que permite afirmar, con un número y no con un adjetivo, que el objetivo "
  "de la planificación inicial —que los tableros no calculen al abrirse— está cumplido.",
  bold=True)

p("Dos elementos no salen del almacén, y es deliberado: el carrito abandonado del tablero "
  "T-1 y el sobre-stock del presente del T-2. Los pide la pantalla con una segunda llamada a "
  "informes simples de PostgreSQL, porque son preguntas del presente —los topes de stock "
  "mínimo y máximo son de hoy y no están en la serie histórica—. El efecto útil es que esos "
  "dos elementos siguen vivos aunque ClickHouse se apague.")

h(2, "4.3 Los dos modelos")

p("El nivel estratégico incluye dos modelos, y los dos se publican con sus métricas reales a "
  "la vista, incluida la que no favorece.")

h(3, "Previsión de demanda")
p("Descomposición estacional con factores encogidos, calculada en el ETL y publicada como "
  "una tabla más del almacén: 510 filas, con tres meses de horizonte para el total, diez "
  "categorías y 159 variantes. Métricas leídas hoy de la propia respuesta del sistema:")
bullet("del backtest: 8,78 %, frente al 12,22 % del método ingenuo estacional que le "
       "sirve de vara. El modelo gana.", bold_prefix="Error medio porcentual ")
bullet("88,46 unidades, con una banda del 80 % de ±334 unidades sobre una previsión de "
       "1.032,89 unidades para agosto de 2026.", bold_prefix="Error absoluto medio: ")
bullet("19 meses. Cada mes del calendario tiene como mucho dos observaciones, y de agosto "
       "a diciembre solo una; la banda lo refleja siendo más ancha.",
       bold_prefix="Historia disponible: ")
p("El sistema declara además, en la propia pantalla, que 1.062 de las 1.221 variantes no "
  "tienen previsión individual por falta de historia y reciben la de su categoría, y que la "
  "historia es un conjunto simulado, de modo que el error mide la calidad del método y no "
  "conocimiento del mercado real de Quevedo.")

h(3, "Alerta de abandono de cliente")
p("Modelo de supervivencia sobre el ritmo de compra de cada cliente. Hoy señala 9 clientes "
  "en alerta de 69 evaluados, 3 de ellos críticos, con $674.285,65 de facturación asociada. "
  "Y publica su propia validación en las tres primeras tarjetas de la pantalla:")
bullet("1,99× sobre el azar.", bold_prefix="Mejora medida: ")
bullet("14 casos positivos sobre 167 evaluaciones.", bold_prefix="Muestra: ")
bullet("«NO · p = 0,1019». Con esa muestra, la mejora sobre el azar no es "
       "estadísticamente significativa.", bold_prefix="Dictamen del propio sistema: ")
p("Que un modelo se publique diciendo que no supera al azar de forma significativa es "
  "deliberado, y es probablemente lo más defendible del nivel estratégico. La causa está en "
  "los datos y el sistema la explica al usuario: en la historia disponible las compras "
  "ocurren a ritmo constante y ningún cliente abandona nunca, así que no hay señal que "
  "aprender. La alerta sirve para priorizar una llamada, no para dar por perdido a un "
  "cliente.")

# ══════════════════════════ 5. EL AGREGADO ══════════════════════════
doc.add_page_break()
h(1, "5. El punto 3 de la consigna: el agregado (AGG)")

h(2, "5.1 Qué se planificó")

p("Al inicio del semestre se creó la carpeta «retailmind/data/agg/», destinada a guardar "
  "agregados precalculados en archivos, de modo que los tableros no tuvieran que calcularlos "
  "al abrirse. Pertenecía al diseño original de tres capas en disco —crudo, intermedio y "
  "agregado—, propio del primer pipeline del proyecto, que extraía a archivos Parquet y "
  "cargaba desde ahí.")

h(2, "5.2 Qué hay hoy en esa carpeta: nada")

p("Se verificó directamente. La carpeta contiene exactamente un archivo, «.gitkeep», de cero "
  "bytes, cuyo único propósito es que una carpeta vacía pueda existir en el control de "
  "versiones. Se creó el 14 de mayo de 2026 y no ha recibido contenido desde entonces.")

p("Más revelador que su vacío es su ausencia de menciones: una búsqueda por todo el "
  "repositorio —código Python, Java y TypeScript, archivos de configuración, composición de "
  "contenedores y documentación— no encuentra ni una sola referencia a «data/agg». La "
  "carpeta nunca llegó a estar conectada a nada. Su hermana «data/stage/» sí conserva el "
  "archivo Parquet de aquel pipeline, que ya no se ejecuta.")

h(2, "5.3 Qué se implementó en su lugar")

p("El agregado se construyó, pero dentro del almacén analítico y no en archivos: son tablas "
  "derivadas de «retailmind_dwh», calculadas por el ETL y refrescadas cada noche por "
  "Airflow. El ETL distingue formalmente tres tipos de tarea, y esa distinción en el código "
  "es la que identifica cuál es un agregado y cuál una copia:")

table(
    ["Tipo de tarea en el ETL", "Qué hace", "Cuántas tablas"],
    [
        ["Tarea de carga",
         "Lee de PostgreSQL y copia al almacén. No es un agregado: es un traslado con "
         "transformación de fila", "18"],
        ["Tarea derivada",
         "Se calcula DENTRO de ClickHouse, con una consulta de inserción sobre otras tablas "
         "del almacén; no vuelve a consultar PostgreSQL", "1"],
        ["Tarea de modelo",
         "También se calcula desde el almacén, pero su transformación es un modelo en "
         "Python que no cabe en una consulta", "2"],
    ],
    widths=[1.7, 3.6, 1.4])

p("Los tres agregados reales del sistema son, por tanto, estos:")

table(
    ["Tabla agregada", "Cómo se calcula", "Se deriva de", "Filas hoy"],
    [
        ["fact_stock_mensual",
         "Cuatro pasos dentro de ClickHouse: saldo de cierre de cada trío (variante, "
         "bodega, mes) con movimiento; malla de meses sin huecos desde el primer movimiento "
         "de cada posición; arrastre del saldo por ventana; y valorización",
         "fact_movimiento_inventario y dim_producto", "22.528"],
        ["fact_prevision_demanda",
         "Descomposición estacional multiplicativa con encogimiento empírico, backtest de "
         "origen móvil y banda de predicción, en Python",
         "fact_venta_linea, dim_producto y dim_fecha", "510"],
        ["fact_alerta_cliente",
         "Supervivencia exponencial con la tasa de compra propia de cada cliente y umbral "
         "de cola, en Python",
         "fact_pedido, dim_cliente, fact_ticket y fact_devolucion", "69"],
    ],
    widths=[1.4, 3.0, 1.6, 0.7])

p("El almacén completo son 21 tablas de modelo con 66.082 filas y 3,36 MiB en disco, más la "
  "bitácora del ETL. Lo refresca el DAG «retailmind_dwh» de Airflow, con 21 tareas de carga "
  "—una por tabla— más una de validación, programado a las 02:00. Su última ejecución, la "
  "noche del 8 de agosto de 2026, completó las 21 tablas con resultado «éxito»; la más lenta "
  "tardó 0,7 segundos.")

h(2, "5.4 ¿Se cumplió el objetivo de la planificación?")

p("Sí, y se puede afirmar con números medidos hoy. El objetivo declarado era que los "
  "tableros no tuvieran que calcular al abrirse. Lo medido:")

table(
    ["Medición", "Resultado"],
    [
        ["Los 7 tableros completos, extremo a extremo por la API (tres medidas cada uno)",
         "entre 0,12 y 0,23 segundos"],
        ["Las 43 rutas de informe compuesto, una tras otra",
         "3,35 s en total (78 ms de media)"],
        ["Leer el agregado fact_stock_mensual (serie de 19 meses)", "3–4 milisegundos"],
        ["Reconstruir esa misma serie en vivo desde el kárdex", "29–34 milisegundos"],
        ["Consulta más pesada del sistema: resumen de analítica sobre 2.931.829 eventos",
         "0,573 segundos"],
    ],
    widths=[4.3, 2.4])

p("Leer el agregado es unas ocho veces más rápido que reconstruirlo, pero conviene ser "
  "escrupuloso con lo que esa cifra significa: la diferencia son 26 milisegundos sobre un "
  "tablero que tarda 200. A la escala actual —66.082 filas— el precálculo no es lo que hace "
  "rápidos a los tableros; los hace rápidos el motor columnar, que agrega decenas de miles "
  "de filas en milisegundos. El propio proyecto lo tiene medido en su comparativa "
  "PostgreSQL–ClickHouse: a esta escala PostgreSQL incluso gana, y ClickHouse solo despega "
  "—unas 17 veces— cuando el volumen llega a los millones de filas.")

p("Entonces, ¿para qué sirve el agregado, si no es para la velocidad? Para la CORRECCIÓN. "
  "Reconstruir el saldo de inventario mes a mes desde el kárdex es un cálculo delicado: hay "
  "que respetar la cronología, arrastrar el saldo por los meses sin movimiento y arrancar "
  "cada cadena en cero. Hacerlo una sola vez en el ETL, validarlo posición por posición "
  "contra el inventario vivo —1.406 de 1.406 posiciones, cero diferencias— y publicarlo solo "
  "si cuadra es una garantía que ninguna consulta escrita al vuelo dentro de un tablero "
  "puede ofrecer.")

h(2, "5.5 Por qué la solución adoptada es preferible a la carpeta de archivos")

p("La comparación no es «archivos contra tablas» en abstracto; es entre lo que cada opción "
  "obliga a hacer para no mentir:")

bullet("una carpeta de archivos precalculados sería una TERCERA fuente de verdad, además "
       "de PostgreSQL y del almacén. Nada garantizaría que un archivo escrito una vez "
       "siguiera coincidiendo con la operación tres semanas después, y el sistema no "
       "tendría cómo avisarlo: un archivo obsoleto se lee igual de rápido que uno "
       "correcto.", bold_prefix="Sincronización: ")
bullet("las tablas derivadas se reconstruyen enteras en cada corrida y se comparan contra "
       "el origen antes de publicarse. Si la validación falla, el intercambio no ocurre y "
       "queda publicada la versión anterior: el sistema nunca muestra una cifra que no "
       "haya cuadrado. Un archivo en disco no tiene ese mecanismo salvo que se programe "
       "aparte.", bold_prefix="Validación: ")
bullet("un archivo hay que leerlo, interpretarlo y filtrarlo en el backend. Una tabla se "
       "consulta con los mismos filtros de período, categoría y bodega que el resto del "
       "tablero, y participa en uniones con las demás tablas del almacén. El tablero T-7 "
       "cruza las 19 tablas del modelo: con archivos, eso habría que programarlo entero.",
       bold_prefix="Consulta: ")
bullet("las tablas del almacén llevan la trazabilidad puesta. La bitácora del ETL registra "
       "cuándo se cargó cada una, cuántas filas escribió y con qué resultado, y el tablero "
       "de gobierno del dato lo muestra. La frescura de un archivo sería su fecha de "
       "modificación.", bold_prefix="Trazabilidad: ")
bullet("todo el nivel estratégico se degrada de forma coherente si el almacén no responde: "
       "los tableros contestan con aviso en lugar de romperse. Una carpeta de archivos "
       "habría creado una cuarta ruta de fallo, distinta de las demás.",
       bold_prefix="Degradación: ")

h(2, "5.6 Veredicto sobre el uso del agregado")

p("El agregado se está usando bien. Se revisó específicamente si había cálculo en vivo donde "
  "debería haber precálculo, o precálculo disponible pero ignorado, y no se encontró:")

bullet("El tablero de rentabilidad usa fact_stock_mensual para la serie de capital y para "
       "la matriz de rotación, y acude al kárdex solo para un dato que el agregado no "
       "contiene: los días transcurridos desde la última venta de cada producto. No hay "
       "reconstrucción duplicada.")
bullet("Los informes de capital inmovilizado y de rotación leen el agregado, no el kárdex.")
bullet("Los dos modelos se calculan en el ETL y las pantallas solo leen su tabla, que es "
       "exactamente lo que corresponde: un modelo que no supera su validación no llega a "
       "publicarse, y esa red se la da la carga atómica.")

p("Con una salvedad honesta, que es el reverso de lo anterior: 18 de las 21 tablas del "
  "almacén no son agregados sino copias del origen, y sobre ellas los informes y los "
  "tableros agregan EN VIVO en cada petición. Eso es correcto hoy —a 66.082 filas cuesta "
  "milisegundos— y sería lo primero que habría que revisar si el volumen creciera dos "
  "órdenes de magnitud. No es deuda técnica: es una decisión proporcionada al tamaño actual, "
  "y conviene registrarla como tal para que una revisión futura tenga de dónde partir.")

# ══════════════════════════ 6. PENDIENTE ══════════════════════════
doc.add_page_break()
h(1, "6. Lo que queda pendiente")

p("Los tres niveles están construidos y en funcionamiento. Lo que sigue abierto, dicho sin "
  "adornos:")

table(
    ["Asunto", "Estado real", "Impacto"],
    [
        ["La carpeta «data/agg/»", "Vacía y sin ninguna referencia en el código",
         "Ninguno funcional. Queda como vestigio del pipeline original; debería eliminarse "
         "o documentarse como histórica, para que no sugiera un componente que no existe"],
        ["Semana 27 de la serie de eventos",
         "Cargada con 19 filas, frente a las ~108.581 de las otras 27 semanas",
         "Es una carga parcial evidente. Los análisis por semana que la incluyan verán un "
         "hueco; el resto de la analítica no se ve afectada"],
        ["Informe de avance sobre la meta",
         "Devuelve un aviso con los parámetros por defecto",
         "Faltan las metas desde agosto de 2026; el informe funciona en cuanto se fijen"],
        ["Trazabilidad por lote y vencimiento",
         "Tabla «lote» creada y vacía; decisión de alcance tomada y registrada",
         "Sin impacto en el flujo actual; obligaría a tocar recepción, inventario, kárdex "
         "y despacho"],
        ["Borrador de ajuste de inventario",
         "El estado se admite, pero no tiene flujo",
         "Menor: exigiría una tabla de detalle de líneas del ajuste"],
        ["Precálculo más allá de las tres tablas derivadas",
         "No existe, y hoy no hace falta",
         "Revisar si el volumen crece dos órdenes de magnitud (sección 5.6)"],
        ["Roles personalizados",
         "La capacidad está construida; hay 0 roles creados",
         "Ninguno: los 9 roles de grupo cubren la operación"],
    ],
    widths=[1.6, 2.3, 2.8])

# ══════════════════════════ 7. DISCREPANCIAS ══════════════════════════
h(1, "7. Diferencias entre la documentación del proyecto y el sistema real")

p("La documentación interna de RetailMind tiene historial de envejecer, así que se "
  "contrastaron sus cifras contra el sistema. Las diferencias se listan aquí. En todos los "
  "casos el documento anterior describe un estado pasado y el sistema ha seguido creciendo "
  "desde entonces: ninguna es un error de funcionamiento.")

table(
    ["Cifra", "Dice la documentación", "Dice el sistema hoy"],
    [
        ["Eventos de navegación (fact_eventos)", "2.823.245",
         "2.931.829, en 28 semanas cargadas"],
        ["Movimientos de kárdex", "13.287", "13.289"],
        ["Órdenes de compra", "865", "868"],
        ["Devoluciones de cliente", "196", "197"],
        ["Tickets de soporte", "248", "249"],
        ["Novedades de envío", "176", "177"],
        ["Líneas de compra", "2.949", "2.952"],
        ["Filas de fact_stock_mensual", "21.122", "22.528"],
        ["Registros de auditoría", "7.073", "7.540"],
        ["Tablas de PostgreSQL (catálogo táctico, v5)", "110", "111"],
        ["Informes simples construidos (catálogo táctico, v5)", "29 de 30", "30 de 30"],
        ["Funciones con privilegio elevado", "13 (bitácora de migración)", "17"],
        ["Concesiones de permisos", "1.355",
         "1.355 a los 9 roles de grupo —la cifra es correcta—, pero 2.192 en total "
         "contando propietario y usuarios técnicos"],
    ],
    widths=[2.3, 1.9, 2.5])

p("Se confirmaron sin cambio, entre otras: 111 tablas, 4.083 pedidos, 10.384 líneas de "
  "venta, 1.406 posiciones de inventario, 95 políticas de seguridad por fila sobre 50 "
  "tablas, 34 disparadores de horario, 9 roles de grupo, 109 columnas con permiso propio en "
  "14 tablas, 21 tablas de almacén con 66.082 filas, 49 controles de validación, 7 tableros, "
  "19 decisiones de dirección y 39 objetivos tácticos compuestos.", size=10)

doc.add_page_break()
h(1, "Anexo. Cómo se verificó cada cifra")

table(
    ["Fuente consultada", "Qué se obtuvo de ahí"],
    [
        ["PostgreSQL 18.4 del contenedor (puerto 5432), en solo lectura",
         "Conteos de tablas y de negocio, políticas de seguridad por fila, disparadores, "
         "permisos por tabla y por columna, roles, ventanas horarias y funciones"],
        ["ClickHouse del contenedor, en solo lectura",
         "Tablas y filas del almacén, bitácora del ETL, filas y semanas de la serie de "
         "eventos, tamaño en disco y tiempos de consulta comparados"],
        ["API en el puerto 8080, autenticada como administrador",
         "Respuesta y tiempo de los 7 tableros y de las 73 rutas de informe; métricas de "
         "los dos modelos, leídas de la propia respuesta"],
        ["Código fuente del backend, del ETL y del frontend",
         "Rutas por departamento, tablas que lee cada tablero, tipo de tarea de cada tabla "
         "del almacén, guardias de proceso y tareas del DAG"],
        ["Sistema de archivos y control de versiones",
         "Contenido de «data/agg/», fecha de su creación y ausencia de referencias en todo "
         "el repositorio"],
    ],
    widths=[2.6, 4.1])

p("No se ejecutó ninguna escritura sobre ninguna base, ni el verificador del ETL —que deja "
  "una fila en la bitácora—, ni el DAG, ni el generador de semanas de eventos.",
  italic=True, size=10)

SALIDA = "RetailMind_TareaF14_Arquitectura_Por_Nivel.docx"
doc.save(SALIDA)
print("Documento generado:", SALIDA)
