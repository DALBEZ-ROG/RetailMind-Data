# -*- coding: utf-8 -*-
"""
Generador de docs/RetailMind_EV09_Documento_UML.docx
Sistema REAL: tienda PyME con back-office; arquitectura hibrida PostgreSQL
(BD retailmind, nucleo operativo transaccional) + ClickHouse (solo analytics).
8 roles, 8 paquetes operativos, 27 casos de uso (CU-O-01..27).
Backend Spring Boot 3.5 / Java 17, frontend Angular 17.

Imagenes: si el PNG existe en docs/diagramas/png/ se incrusta centrado con su
pie de figura; si no existe, se deja un PLACEHOLDER visible para pegar la imagen.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))          # .../docs
PNG_DIR = os.path.join(BASE, "diagramas", "png")
MAX_W = 6.3  # pulgadas

doc = Document()

# ── Estilos ─────────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)

def _hstyle(name, size, rgb):
    st = doc.styles[name]
    st.font.name = "Calibri"; st.font.size = Pt(size)
    st.font.color.rgb = RGBColor(*rgb); st.font.bold = True

_hstyle("Heading 1", 18, (0x1F, 0x3B, 0x5B))
_hstyle("Heading 2", 14, (0x1F, 0x3B, 0x5B))
_hstyle("Heading 3", 12, (0x2E, 0x5A, 0x88))

# ── Helpers ───────────────────────────────────────────────────────────────
def h(level, text): return doc.add_heading(text, level=level)

def p(text="", bold=False, italic=False, size=None):
    par = doc.add_paragraph(); r = par.add_run(text)
    r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    return par

def bullets(items):
    for it in items: doc.add_paragraph(str(it), style="List Bullet")

def page_break(): doc.add_page_break()

def add_toc():
    par = doc.add_paragraph(); run = par.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "Actualice el campo (clic derecho > Actualizar campos) para generar el indice."
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for e in (f1, instr, f2, t, f3): run._r.append(e)

EMBEDDED, PLACEHOLDERS = [], []

def figura(png_name, caption):
    """Incrusta el PNG si existe; si no, deja un placeholder visible."""
    path = os.path.join(PNG_DIR, png_name)
    if os.path.exists(path):
        par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run()
        try:
            run.add_picture(path, width=Inches(MAX_W))
        except Exception:
            run.add_picture(path)
        EMBEDDED.append(png_name)
    else:
        box = doc.add_paragraph(); box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = box.add_run(f"[ INSERTAR IMAGEN: {png_name} ]")
        r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        hint = doc.add_paragraph(); hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rh = hint.add_run("(pegue aqui el diagrama exportado)")
        rh.italic = True; rh.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
        PLACEHOLDERS.append(png_name)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(caption); rc.italic = True; rc.font.size = Pt(9)
    rc.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# PORTADA (placeholder)
# ════════════════════════════════════════════════════════════════════════════
for _ in range(6): doc.add_paragraph()
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("RetailMind"); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = RGBColor(0x1F,0x3B,0x5B)
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Documento UML — Nivel Operativo (EV09)"); r.bold = True; r.font.size = Pt(16)
st2 = doc.add_paragraph(); st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st2.add_run("Tienda online tipo PyME con back-office · Arquitectura hibrida PostgreSQL + ClickHouse")
r.italic = True; r.font.size = Pt(11)
ph = doc.add_paragraph(); ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ph.add_run("\n\n[ PORTADA — a completar por el equipo: logotipo, autores, asignatura, fecha ]")
r.italic = True; r.font.color.rgb = RGBColor(0x90,0x90,0x90)
page_break()

# ════════════════════════════════════════════════════════════════════════════
# INDICE
# ════════════════════════════════════════════════════════════════════════════
h(1, "Indice")
add_toc()
page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCION
# ════════════════════════════════════════════════════════════════════════════
h(1, "1. Introduccion")
p("Este documento reune los modelos UML del NIVEL OPERATIVO de RetailMind, una plataforma web de "
  "comercio electronico tipo PyME con un back-office administrativo completo. El sistema cubre de "
  "extremo a extremo los dos ciclos de negocio que sostienen la operacion: el abastecimiento "
  "(Procure-to-Pay) y la venta (Order-to-Cash), incluyendo catalogo, inventario, bodegas, "
  "facturacion, logistica y devoluciones.")
p("Arquitectura de datos hibrida (separada por responsabilidad):", bold=True)
bullets([
    "PostgreSQL (BD 'retailmind', ~101 tablas transaccionales) es la fuente de verdad de todo lo "
    "operativo: usuarios, catalogo, proveedores, compras, inventario, pedidos, facturas, envios y "
    "devoluciones. Tiene integridad referencial fisica (FKs), columnas generadas y triggers.",
    "ClickHouse se mantiene exclusivamente para analitica (niveles tactico y estrategico); no "
    "participa del ciclo transaccional operativo.",
])
p("Seguridad: autenticacion JWT STATELESS + RBAC en la capa web y, ademas, seguridad a nivel de "
  "base de datos (roles de grupo grp_*, matriz de privilegios, Row-Level Security y restriccion "
  "horaria). La aplicacion asume el rol de motor por transaccion mediante SET LOCAL ROLE.")
p("El sistema define 8 roles: Administrador, Gerente, Vendedor, Encargado de Compras, Encargado de "
  "Bodega, Encargado de Despacho, Cliente y Analista de Datos; y 8 paquetes operativos con 27 casos "
  "de uso (CU-O-01..27).")
p("Nota: la documentacion de steering del repositorio afirma que 'PostgreSQL fue eliminado'. Esa "
  "afirmacion esta desactualizada respecto del codigo y del esquema reales; este documento refleja "
  "el sistema actual, cuyo nucleo operativo corre integramente sobre PostgreSQL.", italic=True)
p("Diagramas incluidos en este documento:", bold=True)
bullets([
    "Casos de uso por paquete (8 diagramas, uno por paquete operativo).",
    "Diagrama de paquetes general (vista de contexto de los 8 paquetes y sus actores).",
    "Diagrama Entidad-Relacion del modelo de datos transaccional (PostgreSQL).",
    "Diagramas de secuencia de casos representativos: realizar pedido, registrar recepcion de "
    "mercancia y procesar devolucion.",
    "Diagrama de despliegue (arquitectura de contenedores).",
    "Diagrama de componentes (frontend, backend y acceso dual a datos).",
])
page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. DIAGRAMAS DE CASOS DE USO POR PAQUETE
# ════════════════════════════════════════════════════════════════════════════
h(1, "4. Diagramas de casos de uso por paquete")
p("El nivel operativo se organiza en 8 paquetes. Cada diagrama muestra los casos de uso del "
  "paquete y los actores que los ejecutan. Los identificadores CU-O-XX son consistentes con el "
  "documento de especificaciones (RetailMind_EV09_Especificaciones_Operativo).")

PAQUETES = [
 ("4.1", "Seguridad y Control de Acceso", "uc_01_seguridad.png",
  ["Administrador (gestiona usuarios y horarios)",
   "Todos los roles (se autentican y obtienen su JWT)",
   "Sistema (aspecto que asume el rol de motor por transaccion, SET LOCAL ROLE)"],
  ["CU-O-01 Iniciar sesion", "CU-O-02 Renovar sesion", "CU-O-03 Cerrar sesion",
   "CU-O-04 Gestionar usuarios y roles", "CU-O-05 Configurar horarios de acceso"],
  "Cubre la autenticacion por email con JWT, la autorizacion RBAC de las rutas y la administracion "
  "de usuarios y de las ventanas horarias por rol (tabla grupo_horario). Es transversal: todos los "
  "actores pasan por el login. La seguridad se refuerza en la BD con roles de grupo, RLS y triggers "
  "de horario; solo el Administrador administra usuarios y horarios."),

 ("4.2", "Catalogo", "uc_02_catalogo.png",
  ["Publico / visitante (consulta sin autenticacion)",
   "Cliente (consulta y genera eventos)",
   "Administrador (gestiona el catalogo)"],
  ["CU-O-06 Consultar catalogo", "CU-O-07 Gestionar categorias", "CU-O-08 Gestionar marcas",
   "CU-O-09 Gestionar productos", "CU-O-10 Gestionar variantes y atributos"],
  "La consulta del catalogo (productos, categorias, marcas y detalle) es publica; la gestion "
  "(altas, ediciones, activacion/desactivacion de categorias, marcas, productos, variantes y "
  "atributos) es exclusiva del Administrador. El precio de venta reside en la variante "
  "(producto_variante), no en el producto."),

 ("4.3", "Proveedores y Compras", "uc_03_compras.png",
  ["Encargado de Compras (emite ordenes, factura)",
   "Encargado de Bodega (registra recepciones)",
   "Gerente / Administrador (consulta y control)"],
  ["CU-O-11 Emitir orden de compra", "CU-O-12 Consultar ordenes de compra",
   "CU-O-13 Registrar recepcion de mercancia", "CU-O-14 Registrar pago a proveedor"],
  "Realiza el ciclo Procure-to-Pay: emitir ordenes a proveedores, recibir la mercancia "
  "(que sube stock y deja rastro en el kardex) y liquidar la deuda mediante pagos contra las "
  "cuentas por pagar. La recepcion puede ser total o parcial y avanza el estado de la orden."),

 ("4.4", "Inventario y Bodegas", "uc_04_inventario.png",
  ["Encargado de Bodega (ejecuta transferencias)",
   "Gerente / Administrador (lectura y control)",
   "Sistema (StockService, invocado por compras, ventas y devoluciones)"],
  ["CU-O-15 Transferir stock entre bodegas", "CU-O-16 Consultar transferencias",
   "CU-O-17 Consultar existencias"],
  "Controla las existencias por variante y bodega y registra sus movimientos (kardex). El caso "
  "operativo expuesto es la transferencia entre bodegas (una salida en origen y una entrada en "
  "destino). Toda variacion de stock pasa por StockService con bloqueo FOR UPDATE y registro en "
  "movimiento_inventario; no se permite stock negativo."),

 ("4.5", "Ventas y Pedidos", "uc_05_ventas.png",
  ["Vendedor (registra pedidos)",
   "Cliente (consulta sus propios pedidos; aislado por RLS)",
   "Gerente / Administrador (consulta global)"],
  ["CU-O-18 Realizar pedido", "CU-O-19 Consultar pedidos", "CU-O-20 Consultar mis pedidos"],
  "Inicia el ciclo Order-to-Cash: registra pedidos con su detalle, toma snapshot del precio "
  "vigente y descuenta stock directo (salida_venta) al confirmarse, manteniendo el historial de "
  "estados. El Cliente solo ve sus propios pedidos por Row-Level Security (parametro app.cliente_id)."),

 ("4.6", "Facturacion", "uc_06_facturacion.png",
  ["Encargado de Compras (factura de compra)",
   "Vendedor (factura de venta)",
   "Gerente / Administrador"],
  ["CU-O-21 Registrar factura de compra", "CU-O-22 Emitir factura de venta",
   "CU-O-23 Generar PDF de factura"],
  "Emite los comprobantes de ambas puntas del negocio: la factura de compra (que abre la cuenta "
  "por pagar al proveedor) y la factura de venta al cliente (con snapshot de datos fiscales), "
  "ambas con su PDF imprimible (iText). Los numeros los genera el sistema y los totales los "
  "recalculan triggers de la BD."),

 ("4.7", "Logistica y Envios", "uc_07_logistica.png",
  ["Encargado de Despacho (despacha y da seguimiento)",
   "Cliente (consulta el estado de su envio)",
   "Vendedor / Gerente / Administrador (segun RBAC)"],
  ["CU-O-24 Despachar pedido", "CU-O-25 Consultar seguimiento de envio"],
  "Despacha los pedidos generando el envio con su guia, el detalle de lo enviado y el primer "
  "evento de seguimiento, avanzando el pedido a 'despachado'. Permite consultar el rastreo del "
  "envio. La direccion se resuelve de la direccion predeterminada del cliente o 'Retiro en tienda'."),

 ("4.8", "Devoluciones", "uc_08_devoluciones.png",
  ["Vendedor / Encargado de Despacho (gestiona la devolucion)",
   "Roles operativos (consulta)"],
  ["CU-O-26 Procesar devolucion", "CU-O-27 Consultar devolucion"],
  "Procesa devoluciones (RMA) de pedidos: valida el motivo, registra los items devueltos con su "
  "estado y accion, reingresa el stock al inventario (entrada_devolucion_cliente) y deja el pedido "
  "en estado 'devuelto'. Impide devolver mas unidades de las compradas."),
]

for num, titulo, png, actores, casos, desc in PAQUETES:
    h(2, f"{num} {titulo}")
    figura(png, f"Figura — Casos de uso del paquete {titulo} ({png})")
    p("Actores:", bold=True)
    bullets(actores)
    p("Casos de uso:", bold=True)
    bullets(casos)
    p("Descripcion:", bold=True)
    p(desc, italic=True)
    page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. DIAGRAMA DE PAQUETES GENERAL
# ════════════════════════════════════════════════════════════════════════════
h(1, "5. Diagrama de paquetes general")
figura("paquetes_general.png", "Figura — Vista de contexto: 8 paquetes operativos y actores (paquetes_general.png)")
p("Descripcion:", bold=True)
p("Vista de contexto del nivel operativo. Presenta los 8 paquetes (Seguridad y Control de Acceso, "
  "Catalogo, Proveedores y Compras, Inventario y Bodegas, Ventas y Pedidos, Facturacion, Logistica "
  "y Envios, Devoluciones) y los 8 actores, con sus usos principales. Todos los actores se "
  "autentican a traves de 'Seguridad y Control de Acceso' (login JWT); el Administrador tiene "
  "acceso total. Los paquetes se agrupan por los dos ciclos de negocio: Procure-to-Pay (Compras, "
  "Inventario, Facturacion de compra) y Order-to-Cash (Ventas, Facturacion de venta, Logistica, "
  "Devoluciones), con Catalogo y Seguridad como soporte transversal.", italic=True)
page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. DIAGRAMA ENTIDAD-RELACION
# ════════════════════════════════════════════════════════════════════════════
h(1, "6. Diagrama Entidad-Relacion (modelo de datos)")
figura("er_general.png", "Figura — Modelo Entidad-Relacion transaccional en PostgreSQL (er_general.png)")
p("Descripcion:", bold=True)
p("A diferencia de la version anterior del sistema (que operaba sobre un esquema analitico), el "
  "modelo operativo actual reside en PostgreSQL con integridad referencial fisica: las relaciones "
  "se materializan con claves foraneas (FOREIGN KEY) reales, no con referencias logicas.", italic=True)
p("Entidades y relaciones principales del modelo transaccional:", bold=True)
bullets([
    "Seguridad: usuario, rol, usuario_rol (N:M), grupo_horario; cliente referencia a usuario.",
    "Catalogo: producto 1:N producto_variante; producto N:M categoria (producto_categoria) y "
    "N:1 marca; atributo 1:N valor_atributo y variante_valor_atributo (N:M con variantes).",
    "Compras: proveedor 1:N orden_compra 1:N orden_compra_detalle; orden 1:N recepcion_mercancia "
    "1:N recepcion_detalle; orden 1:1 factura_compra 1:N factura_compra_detalle; factura_compra "
    "1:1 cuenta_por_pagar 1:N pago_proveedor.",
    "Inventario: inventario (stock por producto_variante + bodega); movimiento_inventario (kardex, "
    "con tipo_movimiento y referencia polimorfica); transferencia_bodega entre bodegas.",
    "Ventas: cliente 1:N pedido 1:N pedido_detalle; pedido N:1 estado_pedido; historial_estado_pedido; "
    "pedido 1:1 factura_venta 1:N factura_venta_detalle.",
    "Logistica: pedido 1:N envio 1:N envio_detalle; envio 1:N seguimiento_envio; transportista, metodo_envio.",
    "Devoluciones: pedido 1:N devolucion 1:N devolucion_detalle; devolucion N:1 motivo_devolucion.",
])
p("Reglas de integridad implementadas en el propio motor:", bold=True)
bullets([
    "Columnas generadas: *_detalle.subtotal = cantidad * precio_unitario es GENERATED ALWAYS "
    "(pedido_detalle, orden_compra_detalle, factura_compra_detalle, factura_venta_detalle).",
    "Totales por trigger: los totales de cabecera (pedido, orden_compra, factura_compra, "
    "factura_venta, devolucion) los recalculan triggers fn_recalcular_total_*; la aplicacion solo lee.",
    "fecha_actualizacion mantenida por triggers *_touch; unicidad de SKU, slugs y numeros de documento.",
    "Row-Level Security en 34 tablas (politicas por rol de grupo y de aislamiento del cliente) y "
    "triggers de horario que bloquean escrituras fuera de la ventana del rol.",
])
page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. DIAGRAMAS DE SECUENCIA
# ════════════════════════════════════════════════════════════════════════════
h(1, "7. Diagramas de secuencia")
p("Interaccion de casos operativos representativos. En todos, la primera sentencia de la "
  "transaccion es 'SET LOCAL ROLE grp_*' (aspecto PgSessionRoleAspect), por lo que el acceso a "
  "datos corre con los privilegios, RLS y horario del rol; al COMMIT el rol se descarta.")

h(2, "7.1 Realizar pedido (Order-to-Cash)")
figura("seq_realizar_pedido.png", "Figura — Secuencia: realizar pedido (seq_realizar_pedido.png)")
p("Descripcion:", bold=True)
p("Flujo: Frontend (Angular) -> VentasController POST /api/ventas/pedidos -> VentasService "
  "(@Transactional). Al abrir la transaccion, PgSessionRoleAspect ejecuta SET LOCAL ROLE grp_* "
  "(y app.cliente_id si es Cliente). VentasService inserta la cabecera del pedido (estado "
  "'confirmado'), y por cada item: consulta la variante activa y su precio (snapshot), inserta el "
  "detalle (el subtotal lo calcula la columna generada) e invoca StockService.mover('salida_venta') "
  "que hace upsert de inventario, SELECT ... FOR UPDATE, inserta el kardex en movimiento_inventario "
  "y actualiza el stock. Un trigger recalcula el total del pedido. Se registra el historial y se "
  "devuelve el pedido con sus totales ya calculados por la BD.", italic=True)

h(2, "7.2 Registrar recepcion de mercancia (actualizacion de stock + kardex)")
figura("seq_recepcion_mercancia.png", "Figura — Secuencia: registrar recepcion de mercancia (seq_recepcion_mercancia.png)")
p("Descripcion:", bold=True)
p("Flujo: Frontend -> ComprasController POST /api/compras/ordenes/{id}/recepciones -> "
  "ComprasService (@Transactional, SET LOCAL ROLE grp_bodega/grp_compras). El servicio valida que "
  "la orden no este 'recibida' ni 'cancelada', crea la cabecera recepcion_mercancia y, por cada "
  "linea: valida que la cantidad recibida no supere lo pendiente, inserta recepcion_detalle, "
  "asegura la fila de inventario (upsert), la bloquea con FOR UPDATE, inserta el kardex "
  "('entrada_compra') en movimiento_inventario, actualiza el stock y suma cantidad_recibida en la "
  "linea de la orden. Finalmente marca la orden como 'recibida' o 'recibida_parcial'. Todo es "
  "atomico: si una linea falla, se revierte la recepcion completa.", italic=True)

h(2, "7.3 Procesar devolucion (reingreso de stock)")
figura("seq_procesar_devolucion.png", "Figura — Secuencia: procesar devolucion (seq_procesar_devolucion.png)")
p("Descripcion:", bold=True)
p("Flujo: Frontend -> VentasController POST /api/ventas/pedidos/{id}/devolucion -> VentasService "
  "(@Transactional, SET LOCAL ROLE). El servicio verifica el estado del pedido (no 'cancelado' ni "
  "'pendiente'), valida el motivo contra motivo_devolucion (activo) y crea la cabecera devolucion. "
  "Por cada item: valida que no se devuelva mas de lo comprado sumando devoluciones previas, "
  "inserta devolucion_detalle e invoca StockService.mover('entrada_devolucion_cliente') que "
  "reingresa el stock con su kardex. Acumula el monto_total, cambia el pedido a 'devuelto' y "
  "registra el historial.", italic=True)
page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. DIAGRAMA DE DESPLIEGUE
# ════════════════════════════════════════════════════════════════════════════
h(1, "8. Diagrama de despliegue")
figura("despliegue.png", "Figura — Arquitectura de despliegue (despliegue.png)")
p("Descripcion:", bold=True)
p("Arquitectura de despliegue orquestada con contenedores. Nodos principales:", italic=True)
bullets([
    "Navegador del usuario: ejecuta la SPA Angular 17 servida como contenido estatico (nginx).",
    "Contenedor Frontend (Angular / nginx): sirve la aplicacion y reenvia las llamadas /api al backend.",
    "Contenedor Backend (Spring Boot 3.5 / Java 17): API REST en el puerto 8080; valida JWT, aplica "
    "RBAC y asume el rol de motor por transaccion (SET LOCAL ROLE).",
    "Contenedor PostgreSQL: base de datos operativa 'retailmind' (nucleo transaccional con FKs, "
    "RLS, triggers y roles grp_*). Es la fuente de verdad del negocio.",
    "Contenedor ClickHouse: almacen analitico columnar, usado solo para los niveles tactico y "
    "estrategico (analytics de eventos), no para la operacion transaccional.",
])
p("Protocolos: el navegador se comunica con el frontend por HTTP(S); el backend accede a "
  "PostgreSQL (JDBC, pool Hikari como usuario retailmind_app) y a ClickHouse (JDBC) por separado. "
  "CORS restringido al origen del frontend.", italic=True)
page_break()

# ════════════════════════════════════════════════════════════════════════════
# 9. DIAGRAMA DE COMPONENTES
# ════════════════════════════════════════════════════════════════════════════
h(1, "9. Diagrama de componentes")
figura("componentes.png", "Figura — Vista de componentes (componentes.png)")
p("Descripcion:", bold=True)
p("Vista de componentes del sistema, en tres capas:", italic=True)
bullets([
    "Frontend Angular (standalone, lazy-loaded) organizado por features: features/operativo "
    "(catalogo, compras, inventario, ventas, horarios), ademas de login, shop, pedidos, perfil y "
    "analytics. La capa core agrupa services (HttpClient), guards, interceptor JWT y models.",
    "Backend Spring Boot por paquetes de dominio: auth (login JWT, PostgresUserRepository), "
    "security (SecurityConfig con RBAC, JwtAuthenticationFilter, PgSessionRoleAspect que ejecuta "
    "SET LOCAL ROLE, y la lista blanca DbGroupRole), y los paquetes operativos catalogo, compras, "
    "inventario, ventas, pedidos, referencias, admin (catalogo/horarios/usuarios) y pdf.",
    "Componente transversal StockService: centraliza el movimiento de stock + kardex y es "
    "reutilizado por compras, ventas e inventario.",
    "Acceso dual a datos: los paquetes operativos usan JdbcTemplate contra PostgreSQL (pgJdbcTemplate) "
    "con consultas parametrizadas; el paquete analytics usa una conexion independiente a ClickHouse. "
    "El aspecto de rol de motor se aplica a com.retailmind..* excepto com.retailmind.analytics..*.",
])
p("La separacion de responsabilidades garantiza que la operacion transaccional (PostgreSQL, con "
  "seguridad de motor) quede aislada de la analitica (ClickHouse).", italic=True)

# ── Guardar ──────────────────────────────────────────────────────────────────
out_path = os.path.join(BASE, "RetailMind_EV09_Documento_UML.docx")
doc.save(out_path)

h1 = sum(1 for x in doc.paragraphs if x.style.name == "Heading 1")
h2 = sum(1 for x in doc.paragraphs if x.style.name == "Heading 2")
print("OK ->", out_path)
print(f"Secciones H1: {h1} | Subsecciones H2: {h2} | Parrafos: {len(doc.paragraphs)}")
print(f"Imagenes incrustadas ({len(EMBEDDED)}): {EMBEDDED}")
print(f"Placeholders pendientes ({len(PLACEHOLDERS)}): {PLACEHOLDERS}")
